"""
Generate the TruthMachine / Sources / Semantics .docx report.

Run:
    python scripts/generate_truth_machine_report.py
Produces:
    docs/reports/Climatefacts-TruthMachine-Sources-Semantics-2026-05-25.docx

The text is hand-authored against the current code base. The script
is the carrier; updating the report means editing this file and
re-running it.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("docs/reports/Climatefacts-TruthMachine-Sources-Semantics-2026-05-25.docx")


def h1(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    return p


def h2(doc, text):
    return doc.add_paragraph(text, style="Heading 2")


def h3(doc, text):
    return doc.add_paragraph(text, style="Heading 3")


def p(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def code_line(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return para


def kv_table(doc, header, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            table.rows[r_i].cells[c_i].text = str(val)
    return table


def build():
    doc = Document()

    # ------------------------------------------------------------------
    # Cover
    # ------------------------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Climatefacts.ai")
    r.bold = True
    r.font.size = Pt(28)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Sources, Semantics, and the Truth-Machine Question")
    r.bold = True
    r.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "How the platform scores sources, builds its data and semantic layers, "
        "indexes claims, and what happens — step by step — when a real user opens "
        "the news feed, the analyser, the deep-search, or the chat. With an honest "
        "verdict on whether this can be called a “truth machine for climate data” "
        "today, and what it would take to deserve that label."
    ).italic = True

    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stamp.add_run(
        "Snapshot 2026-05-25  ·  Repo head: 6798122 (biome+Köppen map layer)  ·  "
        "Owner: CISU Regen  ·  GCP project: climatenews-495412 (europe-west4)"
    )

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "This report is the companion to "
        "docs/reports/Climatefacts-Architecture-Report-2026-05-24.docx and to "
        "docs/improvementplans/Honest-Gap-Audit-2026-05-25.md. The architecture "
        "report describes the platform's intended shape; the gap audit grades each "
        "claim against the running code. This document slices the same system from "
        "a different angle: sources → data → semantics → user actions → birdseye "
        "feasibility. Where claims here disagree with the May-24 report, the gap "
        "audit and the cited code in this document win."
    ).italic = True

    doc.add_page_break()

    # ------------------------------------------------------------------
    # 0. How to read this report
    # ------------------------------------------------------------------
    h1(doc, "0. How to read this report")
    p(
        doc,
        "Each section pairs the same idea twice: once in plain language so a "
        "non-technical reader can follow it, then in code-anchored detail with "
        "file paths and table names so an engineer or an auditor can verify it. "
        "If you only have ten minutes, read sections 1 (executive read), 6 "
        "(architecture verdict), and 7 (truth-machine feasibility). If you want "
        "to understand exactly how a single user click moves through the stack, "
        "read section 5.",
    )
    p(
        doc,
        "Every grade in this report uses the same four labels as the May-25 gap "
        "audit:",
    )
    bullet(doc, "SHIPPED — works end-to-end for a real user, in production today.")
    bullet(doc, "PARTIAL — the scaffolding is real and the code path runs, but the data layer or the UX is incomplete (e.g. an endpoint exists but returns empty, or a verifier runs but produces weak signals).")
    bullet(doc, "STUB — only the shell exists: a route, a component, a column. No working content behind it.")
    bullet(doc, "ASPIRATIONAL — the only evidence is a comment, a docstring, or a copy line. Users see nothing.")

    # ------------------------------------------------------------------
    # 1. Executive read
    # ------------------------------------------------------------------
    h1(doc, "1. Executive read")
    p(
        doc,
        "Climatefacts.ai is built around a single architectural commitment: every "
        "displayed claim should be traceable to the source it came from, and the "
        "verdict on that claim should not be produced by a language model "
        "summarising its own confidence. The platform tries to honour that "
        "commitment through four layers stacked on top of PostgreSQL:",
    )
    bullet(doc, "A source layer — every news outlet, journal, registry, and government dataset that feeds in has (in principle) a credibility tier and three scoring axes: editorial standards, fact-check engagement, transparency.")
    bullet(doc, "A data layer — articles, claims, fact-checks, source profiles, country indicators, projections, and companies live in versioned tables. Every analytical output writes a row to a provenance ledger that records which model, which prompt, which retrieval strategy, and which source articles were used.")
    bullet(doc, "A semantic layer — articles are embedded with OpenAI text-embedding-ada-002 (1536-dim) and indexed with pgvector HNSW. Full-text search runs on a language-aware tsvector index. The two are fused with Reciprocal Rank Fusion plus a shadow knowledge-graph BFS over entities stored as JSONB on claims.")
    bullet(doc, "A verdict layer — claim verdicts come from deterministic Python: numeric grounding (does the number in the claim actually appear in the evidence within 1% tolerance?), SBTi-validation lookup for corporate net-zero claims, ECGT keyword matchers for offset markers. LLMs extract candidate claims; pure code decides truth.")
    p(
        doc,
        "What works today: the agentic chat dispatcher with its 11 cross-checked "
        "skills, the country passport surface, the article search, the hybrid "
        "retrieval pipeline, the migration runner, the rate limiter, the JWT auth "
        "stack, the deterministic numeric-grounding module, and the source-tier "
        "lookup that adds a prior bonus into the credibility math.",
    )
    p(
        doc,
        "What is real but immature: the multi-LLM verifier runs and writes "
        "claims, but real-world coverage is sparse — the average verified-claim "
        "confidence on production articles is around 10/100 and most articles "
        "yield only one verified claim out of a possible five to ten. The "
        "3-axis source scoring schema is in place but its three axes are not yet "
        "used in the credibility calculation; only the legacy single-bonus path "
        "is wired up. The sources page surfaces zero rows for many sources "
        "because the source-profile table is seeded lazily from articles.",
    )
    p(
        doc,
        "What is honestly missing: a knowledge graph as a first-class store "
        "(entities live as JSONB on claims and traversal is BFS), entity "
        "grounding to catch hallucinated names and places, calibration label "
        "volume to make the published curve trustworthy, persona-specific UX "
        "for six of the seven personas described in copy, and auditor-grade "
        "verification rules for the compliance chips that today are mostly "
        "decorative metadata stamps.",
    )
    p(
        doc,
        "The phrase “truth machine for climate data” is the right north star "
        "and the wrong present-tense claim. The architecture deserves the label; "
        "the delivered data does not yet. Section 7 walks through what would "
        "have to be true — scientifically, commercially, for end users, and for "
        "education — before the label is honest.",
    )

    # ------------------------------------------------------------------
    # 2. Source scoring across all four surfaces
    # ------------------------------------------------------------------
    h1(doc, "2. Source scoring across all four surfaces")

    p(
        doc,
        "Source scoring is the single answer to “why should I trust this?” and it "
        "shows up in four places on the platform. Each surface uses the same "
        "underlying tables but exposes a different slice. This section walks each "
        "in plain language first, then in code.",
    )

    # 2.1 The 3-axis schema
    h2(doc, "2.1 The 3-axis schema (the source of truth)")
    p(
        doc,
        "Plain language: every source has three separate ratings on a 0–100 "
        "scale. Editorial standards (does it have a masthead, a corrections "
        "policy, named bylines?), fact-check engagement (is it IFCN-verified, "
        "does it issue retractions when wrong?), and transparency (does it "
        "publish its funding, ownership, methodology, and conflicts of "
        "interest?). Each source also belongs to a tier (T1 to retracted) which "
        "controls the default scores until an editor hand-tunes them.",
    )
    p(
        doc,
        "Code anchor: infrastructure/database/migrations/versions/"
        "041_source_3axis_scoring.sql adds three INTEGER columns "
        "(editorial_score, factcheck_score, transparency_score) to "
        "source_credibility_tiers and seeds them with tier-based defaults. "
        "Excerpt of the tier defaults (matching the migration verbatim):",
    )
    kv_table(
        doc,
        ["Tier", "Editorial", "Fact-check", "Transparency"],
        [
            ["T1 (IFCN-verified, Q1)", "90", "90", "90"],
            ["T2 (mainstream, corrections)", "70", "75", "65"],
            ["T3 (NGO with sourcing)", "55", "50", "60"],
            ["unknown", "30", "25", "30"],
            ["retracted", "5", "5", "5"],
        ],
    )
    p(
        doc,
        "Hand-curated overrides in the same migration: Reuters 92/88/85, "
        "Associated Press 92/90/82, Carbon Brief 88/92/95, IPCC 95/95/98, "
        "Inside Climate News 85/85/80, Google News 40/35/50. Each override "
        "carries a one-line evidence note appended to notes (e.g. "
        "“IFCN-verified, NewsGuard 100/100”).",
    )

    # 2.2 News feed
    h2(doc, "2.2 Surface 1 — News feed and /search")
    p(
        doc,
        "Plain language: when you scan the feed or run a keyword search, every "
        "article carries a green/amber/red credibility chip. That chip is a "
        "stamp applied at ingestion time, not recalculated on the fly. The "
        "stamp lives on the article row itself, derived from the source the "
        "article came from.",
    )
    p(
        doc,
        "Code path: src/backend/app/domains/content/router.py exposes "
        "GET /api/v2/articles; it delegates to ArticleService.search_articles → "
        "ArticleRepository.list_articles. The repository reads "
        "articles.overall_credibility (high / medium / low enum) and "
        "articles.reliability_score (integer 0–100). The frontend ArticleCard "
        "component reads reliability_score with a fallback to "
        "source_credibility_score and renders the colour chip.",
    )
    p(doc, "What does NOT happen here today:", bold=True)
    bullet(doc, "The three axes from migration 041 are not surfaced in the feed — only the single rolled-up reliability_score is shown.")
    bullet(doc, "There is no live recomputation: if a source's editor score is downgraded tomorrow, existing article rows keep their stamped chip until they are re-ingested or backfilled.")

    # 2.3 URL analysis
    h2(doc, "2.3 Surface 2 — URL analysis (/analyze)")
    p(
        doc,
        "Plain language: when you paste a news URL into the analyser, the "
        "platform pulls the article, extracts atomic claims with a language "
        "model, runs a second language model in parallel to corroborate them, "
        "checks the numeric claims against the cited evidence, looks up the "
        "source tier to add a credibility prior, and combines all of that into "
        "an overall score. The verdict on each claim itself is decided by "
        "deterministic Python, not by a language model.",
    )
    p(
        doc,
        "Code path: src/backend/app/services/url_analyzer.py validates the URL "
        "(HTTPS, no localhost), hashes it for a 24-hour dedup cache, writes a "
        "row to url_analyses, then runs the ClaimExtractor and the "
        "multi_llm_verifier. The credibility math lives in "
        "src/backend/app/domains/intelligence/bayesian_credibility.py — "
        "explicitly NOT a Bayesian conjugate update, a weighted-average model "
        "with prior weight 0.3 and evidence weight 0.7. The prior is sourced "
        "from source_credibility_tiers via "
        "app.domains.trust.source_tier_service.get_source_tier_prior(), which "
        "returns (bonus, tier). Bonuses: T1 +30, T2 +15, T3 +5, unknown 0, "
        "retracted −30. A legacy 8-publisher whitelist (Nature, Science, "
        "Elsevier, Springer, Wiley, PLOS, Frontiers, Copernicus) is used as a "
        "fallback when the tier table has not been seeded.",
    )
    p(doc, "Honest gap on this surface (from gap audit + code reading):", bold=True)
    bullet(doc, "The three axes (editorial / factcheck / transparency) are attached to source profiles by source_profiles._attach_credibility_tiers — but they are NOT consumed inside compute_weighted_score(). Only the legacy single prior_bonus is used.")
    bullet(doc, "Average verified-claim confidence on production articles is around 10/100. Most articles yield one verified claim out of a possible five to ten. The pipeline runs; it does not yet produce useful signal at scale.")
    bullet(doc, "Many article URLs in the corpus were synthetic — fake domains and constructed slugs. Migration 040 (purge_synthetic_data) deleted them and a trigger now refuses any insert with is_synthetic=TRUE.")

    # 2.4 Deep search
    h2(doc, "2.4 Surface 3 — Deep search (/deep-search)")
    p(
        doc,
        "Plain language: deep-search runs three things in parallel — your "
        "question against the internal article corpus by semantic similarity, "
        "the same question against Perplexity for external citations, and an "
        "optional weather context if the question mentions one. The answer is "
        "synthesised by a language model and shown with citations. Source "
        "credibility on the internal citations is passed through from the "
        "article row; external Perplexity citations carry no credibility "
        "annotation today.",
    )
    p(
        doc,
        "Code path: src/backend/app/domains/intelligence/deep_search_service.py "
        "in DeepSearchService.search() launches three concurrent tasks: "
        "_search_internal_corpus (pgvector cosine on articles.embedding), "
        "_search_perplexity (external API), _get_weather_context (conditional). "
        "Citations are assembled with reliability_score and overall_credibility "
        "pass-through from each article; nothing is recomputed here.",
    )
    p(doc, "Honest gap:", bold=True)
    bullet(doc, "External web citations (Perplexity) have no credibility tier — they are surfaced raw. A reader sees “Reuters [Tier 1 chip]” next to “some-blog.example.com [no chip]” and the chip's absence is silent.")
    bullet(doc, "The compare mode (/deep-search?mode=compare) is partially implemented per the gap audit; chat follow-up errors had to be surfaced honestly with commit 20a2441 because the previous behaviour swallowed real backend failures into a generic message.")

    # 2.5 Sources page
    h2(doc, "2.5 Surface 4 — Sources page and /api/v2/sources")
    p(
        doc,
        "Plain language: the sources page is meant to be the single place where "
        "you can see every outlet the platform ingests from, with all three "
        "scoring axes visible side by side. Today the endpoint exists and the "
        "schema is right, but the table is seeded lazily from articles. If no "
        "article has been ingested for an outlet yet, that outlet's row is "
        "blank or missing.",
    )
    p(
        doc,
        "Code path: src/backend/app/domains/content/source_profiles.py "
        "list_profiles() queries source_profiles; the SELECT clause projects "
        "source_id, source_name, source_domain, credibility_score, "
        "editorial_standards, fact_check_record, transparency_level, plus "
        "aggregates like total_articles_analyzed and false_claim_rate. If "
        "source_profiles is empty for a domain, _fallback_profiles_from_"
        "articles builds an on-the-fly aggregate from articles. The 3-axis "
        "data from migration 041 is attached separately via "
        "_attach_credibility_tiers, which joins on domain.",
    )
    p(doc, "Honest gap:", bold=True)
    bullet(doc, "Many domain rows have credibility_score = NULL because no articles have been ingested for them yet. The frontend then shows the source as unscored.")
    bullet(doc, "There is no editor UI to update editorial/factcheck/transparency scores — every change today requires a SQL migration. That makes the schema feel more like a registry than a living instrument.")

    # 2.6 Honest summary
    h2(doc, "2.6 Summary — what is robust, what is a stub")
    kv_table(
        doc,
        ["Surface", "Component", "Status", "Honest note"],
        [
            ["Feed / search", "reliability_score chip", "SHIPPED", "Static stamp from ingestion; no recompute on source change."],
            ["Feed / search", "3-axis display", "STUB", "Schema present, never rendered in the feed."],
            ["URL analysis", "Weighted credibility (prior + evidence)", "SHIPPED (code) / PARTIAL (signal)", "Math works; signal weak — avg confidence ~10/100."],
            ["URL analysis", "3-axis integration into score", "STUB", "Axes loaded but not consumed by compute_weighted_score()."],
            ["Deep search", "Internal citation credibility", "SHIPPED", "Pass-through from article row."],
            ["Deep search", "External citation credibility", "STUB", "Perplexity URLs unannotated."],
            ["Sources page", "Tier + 3-axis projection", "PARTIAL", "Endpoint shape correct; data sparse; some sources return 0 rows."],
            ["Sources page", "Editor workflow", "ASPIRATIONAL", "No UI; every change is a SQL migration."],
        ],
    )

    # ------------------------------------------------------------------
    # 3. Data layer
    # ------------------------------------------------------------------
    h1(doc, "3. Data layer — tables, foreign keys, provenance ledger")

    p(
        doc,
        "Plain language: the database is the platform's spine. Articles, "
        "claims, fact-checks, sources, companies, country indicators, and the "
        "provenance ledger are the load-bearing tables. The most important "
        "single design choice is that every analytical output the platform "
        "produces — a claim extraction, a deep-search synthesis, a hallucination "
        "check — writes a row to one ledger table that records exactly how it "
        "was produced. That ledger is what makes the “every verdict traceable” "
        "promise mechanical instead of aspirational.",
    )

    h2(doc, "3.1 Core tables")
    kv_table(
        doc,
        ["Table", "Primary key", "What it stores"],
        [
            ["articles", "article_id (UUID)", "Title, URL, source_name, excerpt, extracted_text, embedding (vector(1536)), search_tsv (tsvector), country_code, published_date, overall_credibility, reliability_score, is_synthetic."],
            ["claims", "claim_id (UUID)", "Atomic propositions extracted from an article. claim_text, claim_type, entities (JSONB), claim_context."],
            ["fact_checks", "id", "One verdict per claim: verification_status (VERIFIED/FALSE/MISLEADING), confidence_score, evidence (JSONB array), climatecheck_hazard mappings."],
            ["claim_provenance", "id (BIGSERIAL)", "Audit ledger: every claim extraction, deep-search synthesis, Cynefin classification, hallucination check, ingestion enrichment writes one row. See §3.3."],
            ["source_credibility_tiers", "source_id", "Per-source tier + the three 0–100 axes added in migration 041 + a one-line evidence note."],
            ["source_profiles", "source_id (UUID)", "Aggregated view: editorial_standards, fact_check_record, transparency_level, total_articles_analyzed, average_reliability_score, false_claim_rate."],
            ["companies", "company_id (UUID)", "10,000+ company rows from SBTi sync. Ticker, name, country_code, sector_nace. Migration 038 dedupes; many remain after partial cleanup."],
            ["company_climate_disclosures", "id", "One row per (company, source, reporting_year). scope1_tco2e, sbti_validated, net_zero_target_year, offset_based_claims."],
            ["country_indicators", "id", "Raw indicators from World Bank, IRENA, OWID, UNFCCC NDC, ND-GAIN, Climate TRACE."],
            ["country_projections", "id", "IPCC AR6 SSP1-2.6 / 2-4.5 / 3-7.0 projections. 180 rows = 20 countries × 3 SSPs × 3 horizons. Atlas backfill pending."],
            ["saved_items", "id", "Polymorphic save table — articles, analyses, searches, companies, feed settings. Migration 042."],
            ["url_analyses", "id", "One row per /analyze submission. url_hash for 24h cache."],
        ],
    )

    h2(doc, "3.2 Indexing strategy")
    p(
        doc,
        "Indexes split into three groups by purpose.",
    )
    p(doc, "Vector index (semantic similarity):", bold=True)
    bullet(doc, "idx_articles_embedding — HNSW (m=16, ef_construction=64), partial WHERE embedding IS NOT NULL, vector_cosine_ops. Replaced an IVFFlat lists=100 index in migration 019 because IVFFlat's recall collapsed past 20k rows and production seeds 50k+.")
    p(doc, "Full-text index (keyword search):", bold=True)
    bullet(doc, "idx_articles_search_tsv — GIN over articles.search_tsv (a generated column produced by clilens_lang_cfg(language_code) — language-aware stemming for en, fi, de, fr, es, it, pt, sv, no, da, nl, ru, tr, hu). Migration 018 fixed an audit finding where hardcoded English stemming mangled ~70% of non-English content.")
    p(doc, "Filter / join indexes (hot paths):", bold=True)
    bullet(doc, "B-tree on articles.published_date, articles.source, articles.overall_credibility.")
    bullet(doc, "B-tree on claims.article_id, claims.claim_type.")
    bullet(doc, "B-tree on fact_checks.verification_status, fact_checks.confidence DESC.")
    bullet(doc, "Partial B-trees on claim_provenance.claim_id, claim_provenance.url_analysis_id, claim_provenance.article_id, and DESC on claim_provenance.created_at.")
    p(
        doc,
        "What is NOT indexed: there is no trigram (pg_trgm) index for fuzzy "
        "name matching, and there is no full-text index on claim_text itself — "
        "claims are filtered by their article membership, not searched "
        "directly.",
    )

    h2(doc, "3.3 The provenance ledger (the truth-machine plumbing)")
    p(
        doc,
        "Plain language: every time the platform writes a claim, a verdict, a "
        "deep-search answer, or a hallucination check, it also writes one row "
        "to claim_provenance recording exactly what model produced the output, "
        "which prompt and prompt-version were used (with a SHA-256 "
        "fingerprint), what retrieval strategy gathered the evidence, which "
        "specific source articles fed in, and what the hallucination and "
        "confidence scores came out as. If a user or an auditor asks “where "
        "did this number come from?”, the answer walks from the displayed "
        "value through this ledger back to the article, back to the source, "
        "back to the upstream registry.",
    )
    p(
        doc,
        "Code anchors:",
    )
    bullet(doc, "Schema: infrastructure/database/migrations/versions/021_claim_provenance.sql. CHECK constraint requires at least one of {claim_id, url_analysis_id, article_id, deep_search_session_id, cynefin_classification_id} to be set.")
    bullet(doc, "Write path: src/backend/app/domains/intelligence/provenance.py — record_provenance() inserts a ProvenanceRecord. extraction_method is one of 'url_analysis_claim_extraction', 'deep_search_synthesis', 'cynefin_classification', 'hallucination_check', 'article_ingestion_enrichment'.")
    bullet(doc, "Negative findings: record_negative_finding() writes a row with confidence=0.0 when the platform looked for something and did not find it — for example, a numeric grounding check that returned false. This is structurally important because it stops “absence of finding” from being indistinguishable from “did not run”.")
    bullet(doc, "Read path: surfaced via /api/methodology/audit-trail/{claim_id} for users and auditors.")
    bullet(doc, "Failure mode: best-effort. If the provenance write fails, the calling pipeline still proceeds — provenance.py logs a warning and returns None. This is the right call (do not break user-facing output on an audit write) but it does mean a small fraction of analytical outputs will lack a ledger entry under load.")

    # ------------------------------------------------------------------
    # 4. Semantic layer
    # ------------------------------------------------------------------
    h1(doc, "4. Semantic layer — embeddings, hybrid retrieval, the shadow KG")

    p(
        doc,
        "Plain language: the platform finds “articles like this question” by "
        "doing three things and combining their answers. It compares the "
        "question to each article by meaning (vector similarity on "
        "embeddings), by keywords (full-text search), and by entity overlap "
        "(walking from the question's named entities to the articles that "
        "mention them). The three ranked lists are merged with a standard "
        "fusion formula, and the top fifteen results are returned. None of "
        "the three is good enough alone; the combination is what makes the "
        "retrieval feel sharp.",
    )

    h2(doc, "4.1 Embeddings (the meaning channel)")
    p(
        doc,
        "Code path: src/backend/app/domains/intelligence/embedding_service.py "
        "populate_embedding() vectorises title + excerpt + extracted_text "
        "(truncated to 32k characters) using OpenAI text-embedding-ada-002 "
        "(1536-dim). Stored in articles.embedding (vector(1536)). Cost is "
        "roughly $0.02 per million tokens; ada-002 is the cheap model, not "
        "the highest-quality model OpenAI sells. There is a near-term "
        "argument to move to text-embedding-3-large for the claim corpus.",
    )
    p(
        doc,
        "Query path: cosine distance via the pgvector <=> operator against the "
        "HNSW index. SET LOCAL hnsw.ef_search = 40 for chat retrieval, 100 "
        "for deep search. ef_search controls recall-vs-latency at query time.",
    )

    h2(doc, "4.2 Full-text search (the keyword channel)")
    p(
        doc,
        "Code path: articles.search_tsv is a generated column produced by "
        "clilens_lang_cfg(language_code) (a SQL function that maps ISO 639-1 "
        "codes to PostgreSQL text-search configurations). Querying uses "
        "websearch_to_tsquery('simple', q) for cross-language matching, or "
        "the per-language stemmer when the caller passes a language. The GIN "
        "index is partial (WHERE search_tsv IS NOT NULL) so unenriched "
        "articles do not bloat it.",
    )

    h2(doc, "4.3 The shadow knowledge graph (the entity channel)")
    p(
        doc,
        "Plain language: there is NO dedicated knowledge-graph table. "
        "Entities — people, organisations, places, policies, events, "
        "technologies, emission sources, concepts — are extracted by a "
        "language model and stored as a JSONB blob on the claim that "
        "mentioned them. When a query needs to find articles that talk about "
        "the same entity, the code does a breadth-first search over the "
        "JSONB column. This works for small queries but it is not a real "
        "graph store. Entity deduplication is per-article; “EU”, “European "
        "Union”, and “Brussels” all live as separate strings on different "
        "claims.",
    )
    p(
        doc,
        "Code path: entity_extraction_service.py runs an LLM prompt that "
        "returns JSON {entities: [{name, type, description}], relationships: "
        "[{source_entity, target_entity, relationship_type, evidence_text}]}. "
        "Entity types include PERSON, ORGANIZATION, LOCATION, POLICY, EVENT, "
        "TECHNOLOGY, EMISSION_SOURCE, CONCEPT. Relationships include CAUSES, "
        "AFFECTS, REGULATES, FUNDS, OPPOSES, MITIGATES, REPORTS_ON, "
        "LOCATED_IN, MEMBER_OF. All stored as JSONB; the GIN index on the "
        "JSONB column is what makes BFS tolerable.",
    )

    h2(doc, "4.4 Hybrid retrieval — how the three channels are combined")
    p(
        doc,
        "Code path: src/backend/app/domains/intelligence/hybrid_rag_service.py "
        "HybridRAGService.retrieve() runs _semantic_search, _fulltext_search, "
        "and _graph_search concurrently, each returning a ranked list. "
        "_reciprocal_rank_fusion() then computes 1 / (k + rank) per list, "
        "sums per document, and returns the top fifteen. RRF is the standard "
        "rank-fusion choice for hybrid IR; it has the property that a "
        "document only needs to be near the top in any one channel to win.",
    )
    p(doc, "Honest summary — what is real semantic search vs what is text matching:", bold=True)
    bullet(doc, "Real semantic: pgvector HNSW cosine on OpenAI ada-002 embeddings. Catches paraphrases, synonyms, concept drift. Weakness: no notion of claim direction (causation, contradiction, support).")
    bullet(doc, "Text matching: tsvector GIN with language-aware stemmers. Precise, fast, offline-capable. Weakness: misses cross-lingual synonymy.")
    bullet(doc, "Shadow KG: JSONB entity overlap walked by BFS. Works for the small queries it sees but cannot do graph-level reasoning, cannot resolve entity duplicates, and cannot answer relational questions.")
    bullet(doc, "Hybrid (the real retrieval): RRF over the three channels. The right design for the platform's size. Hardest gap: there is no semantic understanding of relationships between claims, so a claim that supports another claim and one that contradicts it look similar to the retriever.")

    # ------------------------------------------------------------------
    # 5. User journeys
    # ------------------------------------------------------------------
    h1(doc, "5. User journeys — what runs when you click")

    p(
        doc,
        "Five concrete actions. For each: the route the user is on, the "
        "endpoints that fire, the backend code that runs, the tables that are "
        "read and written, the language model calls (if any), and the honest "
        "status of the path end-to-end.",
    )

    # 5.1 Search
    h2(doc, "5.1 “I open /search and type ‘Norway oil claims’”")
    p(doc, "What you see:", bold=True)
    p(doc, "A list of articles, each with a coloured credibility chip, sorted by relevance and freshness. Filters in the URL for country, credibility tier, date range, tags, category.")
    p(doc, "What runs:", bold=True)
    bullet(doc, "Frontend: src/frontend/src/app/search/page.tsx debounces input 300 ms and calls api.getArticles({q, country, credibility, tags, date_from, date_to, category}).")
    bullet(doc, "Backend: GET /api/v2/articles (src/backend/app/domains/content/router.py) → ArticleService.search_articles → ArticleRepository.list_articles (repository.py).")
    bullet(doc, "DB: SELECT against articles with multilingual FTS or LIKE depending on the q shape; filters on tags JSONB, country_code, overall_credibility, published_date.")
    bullet(doc, "LLM calls: none. This path is pure SQL.")
    bullet(doc, "Credibility chip: read directly from articles.overall_credibility (stamped at ingestion time).")
    p(doc, "Status: SHIPPED. The path is fast and stable. The honest caveat is that the chip is a static stamp, so changes to a source's credibility do not propagate to existing articles until they are re-ingested.", italic=True)

    # 5.2 Deep search
    h2(doc, "5.2 “I open /deep-search and ask ‘How is Norway's offshore drilling policy changing under EU pressure?’”")
    p(doc, "What you see:", bold=True)
    p(doc, "A synthesised answer with inline citations to internal articles and external Perplexity sources, optional weather/climate context, and refinement chips that propose follow-up questions.")
    p(doc, "What runs:", bold=True)
    bullet(doc, "Frontend: src/frontend/src/app/deep-search/page.tsx calls api.deepSearch(query, country, includeWeather). URL state persisted via useUrlState so the query and mode survive page reloads.")
    bullet(doc, "Backend: POST /api/deep-search → DeepSearchService.search() (src/backend/app/domains/intelligence/deep_search_service.py).")
    bullet(doc, "Three concurrent tasks: _search_internal_corpus (pgvector cosine on articles.embedding, ef_search=100), _search_perplexity (external API), _get_weather_context (conditional on weather keywords in the query).")
    bullet(doc, "LLM calls: Perplexity for external search, Anthropic Claude for the final synthesis prompt. The synthesis prompt fingerprint is written to claim_provenance with retrieval_strategy='internal_corpus+perplexity_external+weather'.")
    bullet(doc, "DB writes: claim_provenance row per synthesis. No new article or claim rows are written by deep-search itself.")
    p(doc, "Status: PARTIAL. The path runs; chat follow-up errors are now surfaced honestly after commit 20a2441 (the previous behaviour swallowed real errors into a generic message). Compare mode is implemented but rough. External citations have no credibility annotation.", italic=True)

    # 5.3 URL analysis
    h2(doc, "5.3 “I paste a news URL into /analyze”")
    p(doc, "What you see:", bold=True)
    p(doc, "A breakdown of claims extracted from the article, each with a verification status (VERIFIED / FALSE / MISLEADING / UNVERIFIABLE), a confidence score, and the source articles or evidence URLs that backed the verdict. An overall credibility score and a calibration label.")
    p(doc, "What runs:", bold=True)
    bullet(doc, "Frontend: src/frontend/src/app/analyze/page.tsx submits {url} to POST /api/analyze-url.")
    bullet(doc, "Backend pipeline (src/backend/app/services/url_analyzer.py): validate_url() → generate_url_hash() (SHA-256) → check_existing_analysis() (24h cache) → create_analysis_record() → NewsScraperPool fetches and extracts the body → ClaimExtractor uses Anthropic to extract atomic claims → multi_llm_verifier.verify_claims() runs DeepSeek (primary) + Anthropic (secondary) in parallel and computes Jaccard agreement → numeric_grounding.check_numeric_grounding() validates the numbers in each claim against the cited evidence within 1% tolerance.")
    bullet(doc, "Credibility math: bayesian_credibility.compute_weighted_score() combines a source tier prior (0.3 weight) with the mean of per-claim verification confidences (0.7 weight). The source tier prior comes from get_source_tier_prior() against source_credibility_tiers.")
    bullet(doc, "DB writes: url_analyses (status, overall_credibility, reliability_score), claims (one per extracted claim), fact_checks (one per claim with verdict), and claim_provenance entries for the extraction and the synthesis.")
    p(doc, "Status: PARTIAL. The pipeline is sound. The signal is weak — average verified-claim confidence is around 10/100 and most articles yield one verified claim out of a possible five to ten. Multi-claim extraction (A2 in the gap audit) is the highest-leverage fix.", italic=True)

    # 5.4 Companies
    h2(doc, "5.4 “I open /companies/MSFT”")
    p(doc, "What you see:", bold=True)
    p(doc, "A company profile with Scope 1/2/3 emissions, SBTi validation status, net-zero target year, offset_based_claims flagged with an ECGT chip, and a claim-verification analyser that lets you paste a corporate sustainability claim and see how the platform grades it.")
    p(doc, "What runs:", bold=True)
    bullet(doc, "Frontend: src/frontend/src/app/companies/[ticker]/page.tsx. ?view=business URL toggle is persisted via useUrlState; in business view, CSRD/IFRS S2/TCFD chips stamp wherever the underlying data supports them.")
    bullet(doc, "Backend: GET /api/companies/{ticker} → corporate.repository.get_company (handles both UUID and ticker lookups). Reads companies, company_climate_disclosures, company_climate_claims.")
    bullet(doc, "Data sources: SBTi adapter (Google Sheets, 2026 schema) is live. CDP and NZT adapters return clean 200 with data_source_unavailable warnings — CDP retired its anonymous public CSV in 2024, NZT moved to GraphQL.")
    bullet(doc, "Claim analyser: POST /api/companies/{ticker}/analyze (rate-limited, see Phase 8) runs _analyze_claim() which combines an SBTi-validation lookup (deterministic), an ECGT keyword matcher (offset / carbon credit / climate-neutral markers → ECGT Article 4 flag, deterministic), and a numeric grounding check on any quantitative claim. LLMs may extract candidate claims; the verdict itself is pure Python.")
    bullet(doc, "DB writes: company_climate_claims row per analyser submission with the verdict and an evidence reference.")
    p(doc, "Status: PARTIAL. SBTi data path is the cleanest in the platform; CDP/NZT are honest stubs returning warnings. The 10k company rows include ~70% SBTi target-row duplicates pending migration 038 cleanup, and a UUID-cast bug previously broke /api/companies/{ticker} on ticker lookups (fixed in commit 386e57e).", italic=True)

    # 5.5 Chat
    h2(doc, "5.5 “I open the chat panel anywhere and ask ‘Show me Norway's emissions vs Sweden's’”")
    p(doc, "What you see:", bold=True)
    p(doc, "A short answer plus zero to three suggested actions. Auto actions execute immediately (a URL navigation, a filter application). Confirm actions (URL analysis, bookmark, corporate claim verification, calibration label) gate behind a confirmation modal because they consume quota or mutate state.")
    p(doc, "What runs:", bold=True)
    bullet(doc, "Backend: POST /api/chat. The prompt template renders the 11-skill registry verbatim via render_actions_block_for_prompt() (skills.py). The LLM returns a JSON action like {\"skill\": \"start_deep_search\", \"params\": {\"q\": \"Norway emissions vs Sweden\"}}.")
    bullet(doc, "Single source of truth: src/backend/app/domains/intelligence/skills.py SKILLS_REGISTRY dict (11 entries: navigate, analyze_url, apply_search_filters, apply_map_filters, open_methodology_section, open_country, start_deep_search, bookmark_article, start_calibration_label, open_company, verify_corporate_claim). Each is a frozen Skill dataclass with name, description, mode ('auto' | 'confirm'), parameters tuple, target_surfaces tuple.")
    bullet(doc, "Frontend dispatcher: chatActionDispatcher.ts reads its ChatActionType union from the same registry. Mode 'auto' (7 skills) dispatches immediately; mode 'confirm' (4 skills) runs the confirmation modal first.")
    bullet(doc, "Cross-check: tests/api/test_agentic_skill_pin.py asserts SKILLS_REGISTRY, the chat prompt template, and the frontend dispatcher are perfectly aligned. A drift fails CI and blocks merge — the LLM can never suggest an action the dispatcher cannot execute.")
    p(doc, "Status: SHIPPED. This is the cleanest agentic protocol surface in the platform. The drift-prevention via the pin test is the right pattern; it is what makes the 11 skills a real protocol instead of LLM theatre.", italic=True)

    # ------------------------------------------------------------------
    # 6. Architecture verdict
    # ------------------------------------------------------------------
    h1(doc, "6. Architecture verdict — robust, needs development, missing")

    h2(doc, "6.1 Robust")
    bullet(doc, "Domain-driven split (Content / Intelligence / Trust) keeps the verdict layer testable in isolation.")
    bullet(doc, "Migrations runner with cloud-sql-proxy in Cloud Build. 42 versioned migrations, idempotent runner, schema_migrations_applied tracker.")
    bullet(doc, "Provenance ledger (claim_provenance) — the structural payoff for the truth-machine promise; the ledger writes, the read endpoints surface, the failure mode is non-load-bearing.")
    bullet(doc, "Agentic skills protocol — 11 typed skills, frozen dataclass registry, pin test that blocks merge on drift. This is unusual rigour for a chat layer and it is the right pattern.")
    bullet(doc, "Hybrid retrieval — pgvector HNSW + multilingual FTS + JSONB-entity BFS fused by RRF. The right design for the corpus size.")
    bullet(doc, "Numeric grounding module — pure functions, exhaustive unit tests, wired into the verifier. Catches the failure mode where two language models hallucinate the same wrong number.")
    bullet(doc, "JWT auth with rotating refresh tokens + replay detection. Stripe billing scaffold. Quota envelope returning a structured 429 with upgrade CTA.")
    bullet(doc, "Cloud Run with min-instances=1 + 1800 s request timeout so background tasks (adapter sync, AOI poll) complete reliably.")

    h2(doc, "6.2 Needs development")
    bullet(doc, "Multi-claim extraction yield — the verifier produces ~1 claim per article on average. Lift to 3–5 by re-engineering the extractor prompt and the chunking strategy. Scope estimate from the production-review response: 2–3 days.")
    bullet(doc, "3-axis source scoring integration — migration 041 ships the schema; compute_weighted_score() still uses only the legacy single bonus. Two-day refactor wires the three axes into the math and surfaces them in the feed and the sources page.")
    bullet(doc, "Sources page coverage — /api/v2/sources returns sparse data because source_profiles is seeded lazily. A nightly batch that backfills profiles for every source in the RSS registry would fix this in a day.")
    bullet(doc, "Entity grounding (B4 in the gap audit) — add spaCy NER + cross-LLM entity-overlap check so the platform catches hallucinated names and places, not just hallucinated numbers. 2–3 days.")
    bullet(doc, "External citation credibility — annotate Perplexity URLs with at least a domain tier so deep-search citations have visible provenance.")
    bullet(doc, "Calibration label volume — the curve published on /methodology only stabilises with hundreds to thousands of labels; today the labelled corpus is small.")
    bullet(doc, "Article body completeness — articles store excerpts and partial extracts, not full bodies. Multi-claim extraction yield is partly bottlenecked here.")

    h2(doc, "6.3 Missing")
    bullet(doc, "A first-class knowledge graph store. Today the KG is a JSONB shadow on claims walked by BFS. A real graph (entity_links, entity_resolver tables, canonicalised entity ids with deduplication) would unlock relational queries and graph-aware retrieval.")
    bullet(doc, "Persona-specific UX for six of the seven personas. The architecture report names Consumer / Journalist / ESG / Researcher / Policymaker / Analyst / Business decision-maker. Only the Business view toggle (?view=business) is a delivered persona surface today.")
    bullet(doc, "Auditor-grade compliance verification. CSRD / IFRS S2 / TCFD / TNFD chips are metadata stamps; none of them is paired with an independent verification rule auditable against the underlying directive.")
    bullet(doc, "Document upload analysis. /analyze takes a URL; PDF / Word corporate sustainability reports are not accepted today.")
    bullet(doc, "Research feed. Only the news feed exists; arXiv / Nature / Science RSS with topic filters is in the roadmap (C2).")
    bullet(doc, "Scenario simulation UI. ProjectionsPanel renders static IPCC AR6 cards; an interactive “+1.5°C over 30y for DE” simulator is C3 in the production-review.")
    bullet(doc, "Live editor workflow for source scoring. Today every score change is a SQL migration; there is no admin UI.")
    bullet(doc, "SSE / WebSocket streaming for chat. Answers arrive as one chunk; token-by-token streaming would improve perceived latency considerably.")

    h2(doc, "6.4 Honest data and benchmark gaps")
    p(
        doc,
        "Three places where the underlying data does not yet match the surface "
        "the platform promises:",
    )
    bullet(doc, "Article corpus. After the migration 040 synthetic-data purge, the live article count dropped sharply because most pre-purge rows were synthetic. The platform's article volume has to be rebuilt from real ingestion against the 200+ RSS feeds before any feed-level metric is meaningful.")
    bullet(doc, "Source coverage. The 200+ RSS feeds nominally cover UN-193, but per-country article density is heavily skewed toward English-language sources. A meaningful “source diversity” claim needs balanced ingestion across language and region.")
    bullet(doc, "Reliable benchmarks. There is no external benchmark suite the platform is graded against today. ClimateX, ClimateFEVER, and the IPCC AR6 statement set would all be plausible benchmarks. Building an internal benchmark harness against any of these would convert the verifier's “~10/100 confidence” into an externally-grounded number.")

    # ------------------------------------------------------------------
    # 7. Truth-machine birdseye
    # ------------------------------------------------------------------
    h1(doc, "7. The truth-machine question — birdseye")

    p(
        doc,
        "The phrase “truth machine for climate data” has three load-bearing "
        "words. Truth: a verdict on claims, not a summary of them. Machine: "
        "produced by code paths that an auditor can reproduce, not by a "
        "language model summarising its own confidence. Climate data: the "
        "subject matter is bounded — emissions, projections, regulatory "
        "compliance, climate news. The platform's architecture is built "
        "against all three, but each has to be earned against four "
        "constituencies before the label is honest.",
    )

    h2(doc, "7.1 Scientific feasibility")
    p(
        doc,
        "Plain language: would a working climate scientist trust a number on "
        "this platform enough to cite it?",
    )
    p(doc, "Reliability:", bold=True)
    bullet(doc, "Strong points: provenance ledger, prompt fingerprinting, multi-LLM cross-verification, numeric grounding, IPCC AR6 seed data, source tiers explicitly marking IPCC and Nature as Tier 1.")
    bullet(doc, "Weak points: 1536-dim ada-002 embeddings are dated; ada-002 was deprecated in OpenAI's lineup. Calibration label volume is low. The compliance chips are decorative until paired with verification rules. The shadow KG cannot do relational reasoning.")
    p(doc, "Accessibility:", bold=True)
    bullet(doc, "Methodology page bundles prompts, fingerprints, calibration, drift, hallucination rates. This is more transparency than most science publishers, let alone news platforms, expose.")
    bullet(doc, "Raw indicator endpoints exist (country_indicators, IPCC AR6 projections). A working scientist can pull the numbers behind a passport tile and inspect them.")
    p(doc, "Transparency:", bold=True)
    bullet(doc, "AI provenance JSON-LD on every AI-generated artefact (Schema.org CreativeWork with model + prompt fingerprint + timestamp). Aligns with EU AI Act Article 50 (effective 2 Aug 2026).")
    bullet(doc, "Source-tier rubric is public in the migration file; hand-curated overrides include their evidence note.")
    p(doc, "Honest verdict (scientific): feasible as a working hypothesis; not yet at the trust level a peer-reviewed citation requires. The path to “yes a scientist would cite this” runs through: external benchmark grading, higher-quality embeddings, real entity grounding, and a published calibration curve based on hundreds of labels rather than tens.", italic=True)

    h2(doc, "7.2 Commercial feasibility")
    p(
        doc,
        "Plain language: can this platform be bought by an ESG team, a "
        "newsroom, a research firm, or a corporate risk function and pay its "
        "way?",
    )
    p(doc, "Strong points:", bold=True)
    bullet(doc, "Five-tier freemium grid with structurally defensible upgrade triggers (URL analyses, AOI alerts, embeds, exports, team seats).")
    bullet(doc, "Business view toggle is the right wedge for a B2B buyer — same data, board-ready framing, ECGT and CSRD chips on the same numbers a fiduciary would see.")
    bullet(doc, "Corporate Tracker with SBTi data live, claim analyser with deterministic ECGT keyword matching, projections panel with IPCC AR6 scenarios. Three plausible enterprise sales hooks already exist on the surface.")
    p(doc, "Weak points:", bold=True)
    bullet(doc, "Compliance chips are metadata stamps without verification rules. An enterprise procurement review will surface this within minutes.")
    bullet(doc, "Article corpus is being rebuilt after the synthetic-data purge; feed-level metrics are not yet defensible to a commercial buyer.")
    bullet(doc, "CDP and NZT adapters are honest stubs. For ESG buyers, CDP coverage is table stakes.")
    bullet(doc, "Embeds are routed but production-tested rarely.")
    p(doc, "Honest verdict (commercial): feasible as a pilot product for a friendly buyer (a CISU-aligned newsroom, an academic ESG group, a national policy lab). Not yet feasible as a procurement-ready enterprise SaaS. The first 90 days of commercial sales would have to be consultative — “here's what we ship, here's our roadmap, here's our gap audit” — rather than RFP-driven.", italic=True)

    h2(doc, "7.3 Consumer feasibility")
    p(
        doc,
        "Plain language: would a curious citizen pasting a climate news URL "
        "into the analyser come away with a useful answer they trust?",
    )
    p(doc, "Strong points:", bold=True)
    bullet(doc, "Plain-language interpretation sentences across the country passport and the analyser surface.")
    bullet(doc, "Credibility chips colour-code every article so the reader does not have to read the methodology page first.")
    bullet(doc, "Free tier (3 saved, 3 searches, 2 deep researches) is enough to taste the workflow without sign-up friction.")
    bullet(doc, "Map walkthrough overlay (commit b6363fb) onboards a first-time visitor in seven steps.")
    p(doc, "Weak points:", bold=True)
    bullet(doc, "Verifier yield is low — a citizen pasting a URL today often sees one verified claim with low confidence. That undersells the architecture; until A2 (multi-claim extraction) lifts the yield, the analyser feels weaker than the platform actually is.")
    bullet(doc, "Several UX bugs are tracked in the production-review (Share buttons broken, My Feed save broken, light-on-light contrast in places). All small but cumulatively erode trust.")
    bullet(doc, "Personalisation footprint is one URL toggle. A consumer-first persona surface would help conversion.")
    p(doc, "Honest verdict (consumer): feasible as a credible-looking sandbox today. Reaches “useful enough that a citizen recommends it to a friend” after the verifier yield lifts and the half-dozen UX bugs close. Roughly two weeks of focused work to get there.", italic=True)

    h2(doc, "7.4 Educational feasibility")
    p(
        doc,
        "Plain language: could a journalism school, a sustainability MBA "
        "class, or a high-school climate civics curriculum use the platform "
        "as a teaching instrument?",
    )
    p(doc, "Strong points:", bold=True)
    bullet(doc, "Methodology page is a teaching artefact in itself — prompt registry, calibration curve, drift verdicts, source tiers, corporate verification taxonomy. A class can dissect it.")
    bullet(doc, "Provenance ledger lets students walk from a displayed verdict back to the source articles. That is the right scaffolding for teaching evidence reasoning.")
    bullet(doc, "Country passport's six tabs (Overview / News / Climate Data / Projections / Sources / Claim Ledger) map directly onto a class structure: what is happening, what the press says, what the data says, what the future looks like, who is reporting it, who has been graded on it.")
    bullet(doc, "Free tier gives a class of 30 students enough quota to do an assignment without billing.")
    p(doc, "Weak points:", bold=True)
    bullet(doc, "No exportable lesson plans, no embeddable “teach” blocks. Today a teacher would have to design their own scaffolding around the platform.")
    bullet(doc, "Article body completeness limits how usable individual articles are for close reading.")
    p(doc, "Honest verdict (educational): feasible today as a self-directed teaching instrument and tomorrow (after small additions: lesson plan templates, an /education landing page, downloadable evidence packs) as a curriculum component. This is the constituency where the platform's architectural transparency is the closest to a finished product.", italic=True)

    # ------------------------------------------------------------------
    # 8. One-page summary
    # ------------------------------------------------------------------
    h1(doc, "8. One-page summary")

    p(doc, "What Climatefacts.ai is, in one paragraph:", bold=True)
    p(
        doc,
        "A climate-news fact-checking and intelligence platform that combines "
        "a multi-LLM verifier (with deterministic numeric grounding and "
        "deterministic verdict rules) with a public regulatory data layer "
        "(SBTi live; CDP, NZT stubs), 200+ RSS feeds nominally covering "
        "UN-193, and an 11-skill agentic chat dispatcher whose action surface "
        "is single-sourced from one Python registry and cross-checked by a "
        "CI-blocking pin test. The verdict path is pure Python; LLMs only "
        "extract candidate claims.",
    )

    p(doc, "What it does well today:", bold=True)
    bullet(doc, "Provenance: every analytical output writes to a ledger with model, prompt fingerprint, retrieval strategy, and source article ids.")
    bullet(doc, "Retrieval: HNSW vectors + multilingual FTS + entity BFS fused by RRF. Right design for the corpus size.")
    bullet(doc, "Deterministic verdicts: numeric grounding, SBTi lookup, ECGT keyword matching. LLMs never decide truth.")
    bullet(doc, "Agentic protocol: 11 typed skills, frozen registry, pin test that blocks drift at CI time.")
    bullet(doc, "Transparency surface: methodology page bundles prompts, fingerprints, calibration, drift, hallucination rates.")

    p(doc, "What it does NOT do yet:", bold=True)
    bullet(doc, "Render the 3-axis source scores in the feed or use them in the credibility math.")
    bullet(doc, "Catch hallucinated entities (names, places, organisations) — only numbers are grounded today.")
    bullet(doc, "Surface external (Perplexity) citation credibility in deep search.")
    bullet(doc, "Verify CSRD / TCFD / IFRS S2 / TNFD compliance against the directive — chips today are metadata stamps.")
    bullet(doc, "Personalise UX per persona — six of seven personas are roadmap items.")
    bullet(doc, "Accept document uploads, render scenario simulations interactively, or stream chat tokens.")

    p(doc, "Is it a truth machine for climate data?", bold=True)
    p(
        doc,
        "The architecture deserves the label. The delivered data does not "
        "yet. The gap is concrete and small enough to enumerate: lift "
        "multi-claim extraction yield, wire the 3-axis scoring into the "
        "credibility math, add entity grounding, build a benchmark harness, "
        "and convert at least one compliance chip into a real verification "
        "rule. Until then, the honest framing is: a climate-data verification "
        "platform with an unusually defensible provenance layer, a working "
        "agentic protocol, and a clear path to deserving the truth-machine "
        "label within one focused quarter of work.",
    )

    # Footer
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run(
        "End of report — Climatefacts.ai · 2026-05-25 · "
        "Companion docs: Architecture Report 2026-05-24, Honest Gap Audit 2026-05-25, "
        "Production Review Response 2026-05-25, Resume-Here 2026-05-25."
    )
    r.italic = True
    r.font.size = Pt(9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build()
    print(f"Wrote {out}")

"""
Generate the CARF / CYNEPIC mapping .docx report.

Run:
    python scripts/generate_carf_mapping_report.py
Produces:
    docs/reports/Climatefacts-CARF-CYNEPIC-Mapping-2026-05-25.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("docs/reports/Climatefacts-CARF-CYNEPIC-Mapping-2026-05-25.docx")


def h1(doc, text):
    doc.add_paragraph(text, style="Heading 1")


def h2(doc, text):
    doc.add_paragraph(text, style="Heading 2")


def h3(doc, text):
    doc.add_paragraph(text, style="Heading 3")


def p(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def code_line(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def kv_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            table.rows[r_i].cells[c_i].text = str(val)
    return table


def build():
    doc = Document()

    # --------------------------- Cover ---------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Climatefacts.ai")
    r.bold = True
    r.font.size = Pt(28)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("How the platform applies the CARF / CYNEPIC architecture")
    r.bold = True
    r.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "A component-by-component mapping between the upstream research "
        "framework (github.com/eljaplacido/projectcarfcynepic, BSL 1.1) and "
        "what is actually wired into the climatenews production codebase. "
        "Where each CARF component lives in this repo, how it is invoked, "
        "what is shipped vs partial vs missing, and the honest delta against "
        "the upstream's stated capabilities."
    ).italic = True

    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stamp.add_run(
        "Snapshot 2026-05-25  ·  Repo head: 6798122  ·  Owner: CISU Regen  ·  "
        "Upstream: BSL 1.1 → Apache 2.0 on 2030-02-19"
    )

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "This report builds on the 2026-05-18 audit at "
        "docs/audits/2026-05-18-projectcarfcynepic-comparison.md and the "
        "two companion reports already in this folder (Architecture Report "
        "2026-05-24 and TruthMachine/Sources/Semantics 2026-05-25). License "
        "stance is unchanged: CARF components are treated as algorithmic "
        "inspiration and re-implemented in our own modules. No upstream "
        "source files are vendored."
    ).italic = True

    doc.add_page_break()

    # --------------------------- §0 grade legend ---------------------------
    h1(doc, "0. How to read this report")
    p(
        doc,
        "Each CARF/CYNEPIC component is graded against the running climatenews "
        "code using the four-label scale used in the rest of this folder:",
    )
    bullet(doc, "SHIPPED — fully re-implemented in climatenews and serving production traffic today.")
    bullet(doc, "PARTIAL — re-implemented in a thinner form; the climatenews version captures the idea but drops some of the upstream's depth or signal richness.")
    bullet(doc, "STUB — the integration shell exists (a route, a client class, an env var) but the real backing service is not deployed; falls back gracefully when called.")
    bullet(doc, "MISSING — explicitly identified in the 2026-05-18 audit as worth porting; no module in climatenews today.")
    bullet(doc, "DECLINED — evaluated and consciously not ported, with the reason recorded in the audit.")

    p(
        doc,
        "Every component cites the upstream concept first, then the "
        "climatenews file:line that implements it (or the audit note that "
        "explains why it was not ported).",
    )

    # --------------------------- §1 CARF in one paragraph ---------------------------
    h1(doc, "1. CARF / CYNEPIC in one paragraph")
    p(
        doc,
        "CARF is the upstream research framework — Complexity-Adaptive and "
        "Context-Aware Reasoning Fabric. Its central thesis is that modern "
        "AI systems sound confident without disclosing where their confidence "
        "comes from, and the fix is to make every query pass through an "
        "epistemic router that decides which kind of reasoning the question "
        "actually needs. The router uses the Cynefin framework's five "
        "domains. Clear questions go to deterministic rule lookups. "
        "Complicated ones go to causal inference (DoWhy / EconML). Complex "
        "ones go to Bayesian inference (PyMC) with separated epistemic and "
        "aleatoric uncertainty. Chaotic ones trip a circuit breaker. "
        "Ambiguous (Disorder) ones escalate to a human. Around that core, "
        "CARF wraps a Guardian policy layer (YAML + CSL-Core + OPA), an "
        "8-signal hallucination sentinel called H-Neuron, a ChimeraOracle "
        "fast-path of pre-trained causal models, an EpistemicState provenance "
        "schema, three governance modules (MAP, PRICE, RESOLVE), and a "
        "monitoring stack (drift, bias, convergence). The whole system is "
        "graded against 43 falsifiable hypotheses with a published "
        "benchmark report.",
    )
    p(
        doc,
        "Climatefacts.ai is one downstream application of that framework, "
        "narrowed to climate news. It ports the pieces of CARF that fit the "
        "climate-news problem, re-implements them cleanly under our own "
        "license, and declines the pieces that would be over-engineering at "
        "our scale. The rest of this report is the per-component evidence "
        "for that statement.",
    )

    # --------------------------- §2 Architectural mapping at a glance ---------------------------
    h1(doc, "2. The mapping at a glance")
    p(
        doc,
        "The full picture in one table. Each row shows an upstream CARF "
        "component, where (if anywhere) it lives in this repo, and the "
        "grade.",
    )

    kv_table(
        doc,
        ["CARF / CYNEPIC component", "Climatefacts.ai implementation", "Grade"],
        [
            ["Cynefin Router (5 domains → 5 engines)", "src/backend/app/domains/intelligence/cynefin_router.py — keyword scoring + LLM fallback, prompt ported from upstream config/prompts.yaml#router", "SHIPPED (labelling only — routing to engines is partial)"],
            ["CARF HTTP backend (FastAPI, /query, /causal, /counterfactual, …)", "src/backend/app/domains/intelligence/carf_integration.py — httpx client to CARF_API_URL, graceful fallback to CynefinRouter", "STUB (client deployed; backend not deployed)"],
            ["CARF API surface (90+ endpoints across 17 routers)", "api/carf_routes.py — 9 endpoints under /api/carf/* exposing classify, causal, counterfactual, hallucination, article analysis, entity graph", "PARTIAL (9 of 90+; the high-value ones)"],
            ["Bayesian Inference (PyMC, three-mode, 90% credible interval, epistemic vs aleatoric)", "src/backend/app/domains/intelligence/bayesian_credibility.py — weighted-average (not conjugate Bayesian); single point estimate; no PyMC", "PARTIAL (idea ported, depth dropped)"],
            ["Causal Inference (DoWhy / EconML, refutation battery)", "src/backend/app/domains/intelligence/causal_analysis.py + causal_claim_analyzer.py — LLM-only causal extraction (cause / effect / mechanism / counterfactual JSON)", "PARTIAL (LLM extraction; no statistical refutation)"],
            ["H-Neuron Sentinel (8-signal weighted hallucination fusion)", "src/backend/app/domains/intelligence/hallucination_detector.py — 3 signals (entity overlap 0.3 + statistic accuracy 0.3 + LLM grounding 0.4)", "PARTIAL (3 of 8 signals)"],
            ["EpistemicState Schema (6 analytical layers, aleatoric vs epistemic uncertainty)", "infrastructure/database/migrations/versions/021_claim_provenance.sql + src/backend/app/domains/intelligence/provenance.py — ledger with 5 extraction methods, single confidence field", "PARTIAL (ledger present, uncertainty split missing)"],
            ["Guardian Policy Layer (YAML + CSL-Core + OPA, fail-closed)", "src/backend/app/domains/intelligence/editorial_gate.py — rules hardcoded in Python (PUBLISH ≥60, HOLD 30–59, ESCALATE <30); no YAML, no CSL, no OPA", "PARTIAL (idea ported, externalisation missing)"],
            ["Drift Detector (KL on routing distribution)", "src/backend/app/domains/intelligence/drift_detection.py — KL on source-mix instead; verdict thresholds stable/minor/notable/significant", "PARTIAL (different distribution tracked)"],
            ["Bias Auditor (chi-squared on memory corpus)", "(not implemented)", "MISSING — item #2 in 2026-05-18 audit"],
            ["MAP–PRICE–RESOLVE governance trio", "(not implemented)", "MISSING — item #5 in 2026-05-18 audit"],
            ["ChimeraOracle (pre-trained CausalForestDML fast-path, <100ms)", "(not implemented)", "MISSING — defer until causal pipeline is real"],
            ["Three-Layer RAG (vector + graph + symbolic, RRF)", "src/backend/app/domains/intelligence/hybrid_rag_service.py — pgvector HNSW + tsvector GIN + JSONB-entity BFS, fused by RRF", "SHIPPED (the cleanest match in the platform)"],
            ["EU AI Act compliance reporting", "src/backend/app/domains/intelligence/carf_integration.py check_eu_ai_compliance (proxied) + Schema.org CreativeWork JSON-LD on every AI artefact", "PARTIAL (Article 50 disclosure shipped; full Article 9/12/13/14 reporting proxied)"],
            ["Versioned prompts.yaml (router / causal-analyst / bayesian-explorer / deterministic-runner / policy-check / self-correction)", "src/backend/app/domains/intelligence/prompts.py — _REGISTRY with SHA-256 fingerprints; cynefin_classifier prompt ported; self-correction missing", "PARTIAL"],
            ["Frontend visualisations (CausalDAG, BayesianPanel, CynefinMatrix, SensitivityPlot, RiskTopographyTab, InterventionSimulator)", "src/frontend/src/components/ — ArgumentationGraph, CredibilityGauge, DecomposedConfidenceChart; no DAG, no Cynefin matrix, no sensitivity plot", "MISSING — item #6 in 2026-05-18 audit"],
            ["MCP server (18 cognitive tools)", "(not applicable)", "DECLINED — no MCP consumers in roadmap"],
            ["LangGraph StateGraph workflow", "(not applicable)", "DECLINED — services.py + Celery already orchestrate"],
            ["OPA / Open Policy Agent", "(not applicable)", "DECLINED — YAML policies would suffice if we add policy externalisation"],
            ["HumanLayer approval workflows", "(not applicable)", "DECLINED — defer until moderation queue exists"],
            ["Kafka audit log", "(Postgres audit_events table)", "DECLINED — Postgres sufficient at scale"],
            ["TLA+ formal specs (StateGraph, EscalationProtocol)", "(not applicable)", "DECLINED — not justifiable for our truth-machine grade"],
            ["17 demo scenarios / 43-hypothesis benchmark harness", "(not applicable)", "DECLINED — upstream scenarios are Scope-3 / supply-chain; we will build a climate-news DeepEval suite"],
            ["Firebase Auth / Cloud SQL", "Cloud SQL already + custom JWT auth", "DECLINED — already on GCP Cloud SQL"],
        ],
    )

    # --------------------------- §3 Cynefin Router ---------------------------
    h1(doc, "3. Cynefin Router — the centrepiece")

    p(
        doc,
        "Plain language: every climate question the platform answers gets "
        "tagged with a complexity domain. “What is the capital of France?” "
        "is Clear. “Why does the Gulf Stream weaken?” is Complicated. "
        "“What will Norway's emissions look like under SSP3-7.0 in 2050?” "
        "is Complex. “Wildfire moving toward this town in the next two "
        "hours” is Chaotic. The platform decides which kind of question it "
        "is before it decides how to answer.",
    )

    h2(doc, "3.1 How the upstream does it")
    bullet(doc, "Five domains: Clear, Complicated, Complex, Chaotic, Disorder.")
    bullet(doc, "Each domain has a target engine: rule lookup, DoWhy causal, PyMC Bayesian, circuit breaker, human escalation.")
    bullet(doc, "Router achieves 89.5% F1 (H0 benchmark).")
    bullet(doc, "Entropy-aware routing — a Disorder verdict triggers clarification before any reasoning runs.")

    h2(doc, "3.2 How climatenews does it")
    p(
        doc,
        "File: src/backend/app/domains/intelligence/cynefin_router.py "
        "(315 lines). CynefinRouter.classify() first runs a keyword score "
        "against four lists:",
    )
    bullet(doc, "CLEAR_KEYWORDS: 'what is', 'when did', 'how many', 'define', 'current temperature', 'latest data', 'who is', 'where is', 'list', 'name of', 'which country', 'how much', 'total', 'average', 'what year'.")
    bullet(doc, "COMPLICATED_KEYWORDS: 'how does', 'what causes', 'compare', 'relationship between', 'impact of', 'explain', 'difference between', 'why does', 'mechanism', 'process', 'analysis', 'correlation', 'effect of', 'influence', 'contribute to', 'factors'.")
    bullet(doc, "COMPLEX_KEYWORDS: 'predict', 'forecast', 'scenario', 'what if', 'future', 'trend', 'model', 'tipping point', 'feedback loop', 'systemic', 'cascade', 'long-term', 'projection', 'emerging', 'evolving', 'adaptation', 'uncertainty', 'transition'.")
    bullet(doc, "CHAOTIC_KEYWORDS: 'emergency', 'crisis', 'sudden', 'unprecedented', 'breaking', 'urgent', 'disaster', 'catastrophic', 'extreme event', 'collapse', 'immediate', 'critical', 'alert', 'evacuate'.")
    p(
        doc,
        "Each keyword's weight is 1.0 + 0.3 × word count, so longer phrases "
        "(“tipping point”, “feedback loop”) dominate single tokens. If no "
        "keyword matches, the router falls through to an LLM classification "
        "(_llm_classify, lines 232–315). The LLM prompt is sourced from the "
        "central prompt registry (prompts.py get_prompt('cynefin_classifier')) "
        "so it carries a SHA-256 fingerprint and a version, and the "
        "classification writes a row to claim_provenance via "
        "_maybe_record_provenance (lines 139–176) tagged extraction_method="
        "'cynefin_classification'.",
    )
    p(
        doc,
        "Disorder handling: if the LLM returns raw_domain='disorder', the "
        "router maps it to 'complicated' with confidence capped at 0.4 so "
        "the caller can flag for clarification (lines 285–296). This is the "
        "honest version of disorder routing — without a human escalation "
        "queue we cannot do better.",
    )

    h2(doc, "3.3 The four strategies")
    kv_table(
        doc,
        ["Domain", "STRATEGY_MAP value", "STRATEGY_DESCRIPTIONS value"],
        [
            ["clear", "direct_lookup", "Simple database query for factual answers."],
            ["complicated", "multi_source_analysis", "HybridRAG retrieval with evidence chain analysis."],
            ["complex", "causal_analysis", "Full causal pipeline with counterfactual reasoning."],
            ["chaotic", "rapid_assessment", "Quick summary with high uncertainty flags."],
        ],
    )

    h2(doc, "3.4 Honest gap")
    bullet(doc, "We label Cynefin domains; we do not yet route into different engines based on them. A Clear question goes through the same multi-source analysis pipeline as a Complicated one today. The upstream's hard-edged engine routing is the next step (audit note: “editorial_gate.py could route Clear → cache, Complicated → DoWhy, Complex → Bayesian, Chaotic → human review”).")
    bullet(doc, "No published F1 benchmark for our climate-news domain mix. The upstream's 89.5% F1 was on a different corpus.")
    bullet(doc, "No human escalation queue — Disorder falls back to Complicated with a capped confidence.")

    # --------------------------- §4 CARF HTTP backend ---------------------------
    h1(doc, "4. CARF HTTP backend — the stub")

    p(
        doc,
        "Plain language: there is a client in climatenews that knows how to "
        "talk to the upstream CARF reasoning service over HTTP. If the "
        "service were deployed and configured, the platform would proxy "
        "complexity classification, causal analysis, counterfactuals, "
        "hallucination checks, and EU AI Act compliance reports through it. "
        "In production today the service is not deployed; every call falls "
        "through to local code paths that we built ourselves.",
    )

    h2(doc, "4.1 The client")
    p(
        doc,
        "File: src/backend/app/domains/intelligence/carf_integration.py. "
        "CARFIntegration class with httpx.AsyncClient, configured by three "
        "env vars: CARF_API_URL (default http://localhost:8000), "
        "CARF_API_KEY, CARF_TIMEOUT (default 30s). Methods:",
    )
    bullet(doc, "health_check() — GET /health.")
    bullet(doc, "classify_complexity(text, context) — POST /query with mode='classify_only'. Falls back to CynefinRouter when unreachable.")
    bullet(doc, "analyze_causal_claim(claim, evidence) — POST /query with analysis_type='causal'. Returns {'status': 'carf_unavailable'} when down.")
    bullet(doc, "counterfactual_analysis(scenario, intervention) — POST /query with analysis_type='counterfactual'.")
    bullet(doc, "check_hallucination(generated_text, source_texts) — POST /query with analysis_type='grounding_check'.")
    bullet(doc, "check_eu_ai_compliance(analysis_output) — POST /query with context='eu_ai_act_compliance'. Returns {'compliant': True, 'note': 'CARF compliance check unavailable'} when down — note that this default is structurally optimistic and worth re-examining.")
    bullet(doc, "enhanced_article_analysis(article_id, db) — pulls the article from the articles table, runs classify_complexity + analyze_causal_claim + check_eu_ai_compliance concurrently with asyncio.gather, falls back per branch.")

    h2(doc, "4.2 The API surface")
    p(
        doc,
        "File: api/carf_routes.py. Nine endpoints under /api/carf/* :",
    )
    bullet(doc, "GET /api/carf/status — returns {configured, health, capabilities}.")
    bullet(doc, "POST /api/carf/classify — Cynefin classification.")
    bullet(doc, "POST /api/carf/causal — causal inference (auth-optional).")
    bullet(doc, "POST /api/carf/counterfactual — counterfactual analysis (auth-required).")
    bullet(doc, "POST /api/carf/hallucination-check — quick grounding check.")
    bullet(doc, "POST /api/carf/article-analysis — full enhanced pipeline on a stored article.")
    bullet(doc, "POST /api/carf/analyze-claim — full local CARF-style pipeline (Cynefin → CausalClaimAnalyzer → HallucinationDetector). This is the route that actually runs in production because it does not depend on the CARF backend.")
    bullet(doc, "POST /api/carf/hallucination-check-full — entity + statistic + LLM grounding via HallucinationDetector.")
    bullet(doc, "GET /api/carf/entity-graph/{article_id} — returns entities, entity relationships, and articles connected via shared entities. Reads from article_entities, entities, entity_relationships tables.")

    h2(doc, "4.3 Honest gap")
    bullet(doc, "STUB grade is honest: the HTTP proxy is wired but the upstream service is not deployed alongside climatenews. Every CARF call that requires the remote service returns 'unavailable' and falls through to local fallbacks.")
    bullet(doc, "The 'analyze-claim' route is the one that proves the pattern can run locally — it composes CynefinRouter + CausalClaimAnalyzer + HallucinationDetector directly, no HTTP. That route is the truth-machine pipeline on the local code path.")

    # --------------------------- §5 Bayesian credibility ---------------------------
    h1(doc, "5. Bayesian credibility — borrowed name, weighted-average math")

    p(
        doc,
        "Plain language: when the platform produces a credibility score for "
        "a URL it analysed, it does NOT run a Bayesian inference. The "
        "module is called bayesian_credibility but the math is a weighted "
        "average between a source-tier prior (30% weight) and the average "
        "of per-claim verification confidences (70% weight). The upstream "
        "CARF runs full PyMC posterior updates and separates epistemic "
        "uncertainty (we don't know) from aleatoric uncertainty (the world "
        "is noisy). We do not.",
    )

    h2(doc, "5.1 What climatenews does")
    p(
        doc,
        "File: src/backend/app/domains/intelligence/bayesian_credibility.py. "
        "WeightedCredibilityService.compute_weighted_score() composes:",
    )
    bullet(doc, "Prior: base + venue bonus. Base is 50 for news, 60 for research reports, 55 for policy documents, 40 for preprints. Venue bonus comes from get_source_tier_prior() against source_credibility_tiers (T1 +30, T2 +15, T3 +5, unknown 0, retracted −30). A legacy 8-publisher whitelist (Nature, Science, Elsevier, Springer, Wiley, PLOS, Frontiers, Copernicus) is used as fallback when the tier table has not been seeded.")
    bullet(doc, "Evidence: mean of normalised per-claim fact-check confidences (0–1).")
    bullet(doc, "Posterior: (prior × 0.3) + (evidence_mean × 0.7), bounded to 0–100.")

    h2(doc, "5.2 What is missing vs CARF")
    bullet(doc, "No conjugate-Bayesian update. The 0.3/0.7 weighting is fixed; CARF's PyMC path produces a real posterior distribution.")
    bullet(doc, "No 90% credible interval. We return a single point estimate, not a (lo, hi) band.")
    bullet(doc, "No epistemic vs aleatoric split. CARF reports them separately; we conflate them into a single confidence number.")
    bullet(doc, "No 'will never fabricate posterior updates' honesty constraint. CARF returns priors unchanged when no real data is available; we still apply LLM-derived signals into the evidence term.")

    p(
        doc,
        "The 2026-05-18 audit lists this as item #3 (3 days of work to add a "
        "mode='mcmc' path with PyMC, report a 90% CI, and split the two "
        "uncertainty kinds). It is the highest-leverage Bayesian upgrade "
        "available to us.",
        italic=True,
    )

    # --------------------------- §6 Causal analysis ---------------------------
    h1(doc, "6. Causal analysis — LLM extraction without statistical refutation")

    p(
        doc,
        "Plain language: when the platform analyses a causal claim — say, "
        "“EU carbon taxes reduced manufacturing emissions by 15%” — it asks "
        "an LLM to extract the cause, the effect, the mechanism, and a "
        "counterfactual sentence. It does NOT run the upstream's DoWhy "
        "refutation battery, which would test the claim against placebo "
        "treatments, random common causes, data subsets, and an E-value "
        "sensitivity check. The result is a structured JSON description of "
        "the causal claim, not a falsifiable statistical estimate.",
    )

    h2(doc, "6.1 What climatenews does")
    p(
        doc,
        "File: src/backend/app/domains/intelligence/causal_analysis.py. "
        "CausalAnalysisService uses DeepSeek (configurable via "
        "DEEPSEEK_API_KEY + DEEPSEEK_MODEL env vars, default deepseek-chat) "
        "with a prompt that returns:",
    )
    code_line(doc, "{")
    code_line(doc, "  \"cause\": \"...\",")
    code_line(doc, "  \"effect\": \"...\",")
    code_line(doc, "  \"mechanism\": \"...\",")
    code_line(doc, "  \"confidence\": 0.75,")
    code_line(doc, "  \"counterfactual\": \"...\",")
    code_line(doc, "  \"causal_chain\": [\"Step 1\", \"Step 2\", \"Step 3\"],")
    code_line(doc, "  \"causal_type\": \"direct|indirect|correlational|spurious\",")
    code_line(doc, "  \"strength\": \"strong|moderate|weak\"")
    code_line(doc, "}")
    p(
        doc,
        "The JSON is then stored on the fact_checks row via store_analysis() "
        "which merges into the metadata JSONB column. There is no "
        "intervention modelling, no DAG construction, no placebo test.",
    )

    h2(doc, "6.2 What is missing vs CARF")
    bullet(doc, "DoWhy refutation battery — the upstream runs placebo treatment, random common cause, 80% data-subset, and E-value sensitivity refutations on every causal claim. We run none of these.")
    bullet(doc, "1,138× MSE accuracy ratio over raw LLM (CARF H1 benchmark). Our LLM-only path inherits raw-LLM accuracy on the causal axis.")
    bullet(doc, "ChimeraOracle fast-path — CARF pre-trains 5 CausalForestDML models and serves <100ms causal-effect predictions. We have no pre-trained models on this axis.")

    p(
        doc,
        "The 2026-05-18 audit lists this as item #4 (4 days of work to "
        "build causal_refutation.py and call it from causal_claim_analyzer "
        "after the LLM extracts an edge). On the temperature/CO₂ and IPCC "
        "tabular series the platform already ingests, running DoWhy once "
        "per headline claim would produce one real causal number per "
        "article — citable, falsifiable, and uniquely defensible.",
        italic=True,
    )

    # --------------------------- §7 H-Neuron sentinel ---------------------------
    h1(doc, "7. H-Neuron sentinel — 3 signals of 8")

    p(
        doc,
        "Plain language: when the platform generates text (a deep-search "
        "synthesis, a plain-language interpretation, an article excerpt "
        "enrichment), it runs a hallucination check that compares the "
        "generated text against the source documents it was supposed to be "
        "grounded in. Our check uses three signals; the upstream uses eight.",
    )

    h2(doc, "7.1 What climatenews does")
    p(
        doc,
        "File: src/backend/app/domains/intelligence/hallucination_detector.py. "
        "HallucinationDetector.check() runs three checks and fuses them:",
    )
    bullet(doc, "Entity overlap (weight 0.3) — _check_entity_overlap compares named entities in the generated text against entities in the sources.")
    bullet(doc, "Statistic accuracy (weight 0.3) — _check_statistics tokenises numbers in the generated text and verifies them against numbers in the sources within tolerance.")
    bullet(doc, "LLM grounding (weight 0.4) — _llm_grounding_check asks an LLM to flag unsupported sentences.")
    p(
        doc,
        "Composite: hallucination_risk = 0.3·(1 − entity) + 0.3·(1 − stat) + "
        "0.4·llm_risk. is_grounded = hallucination_risk < 0.4. The result "
        "writes flagged segments + overall confidence + the three component "
        "scores.",
    )

    h2(doc, "7.2 What the upstream's H-Neuron adds")
    p(
        doc,
        "The upstream H-Neuron sentinel runs eight signals with this "
        "weighting (from src/services/h_neuron_interceptor.py in CARF):",
    )
    kv_table(
        doc,
        ["Signal", "Weight", "Caught"],
        [
            ["deepeval_risk", "0.35", "DeepEval-judged hallucination probability"],
            ["confidence_risk", "0.20", "Model-reported confidence too high vs evidence"],
            ["epistemic_uncertainty", "0.15", "We do not know enough to answer"],
            ["reflection_risk", "0.10", "Self-critique mismatch"],
            ["brevity", "0.05", "Output too short for the question"],
            ["shallow_reasoning", "0.05", "No causal/mechanistic content"],
            ["irrelevancy", "0.07", "Off-topic content"],
            ["verbosity", "0.03", "Padded output"],
        ],
    )
    p(
        doc,
        "The 2026-05-18 audit lists this as item #1 (1 day to extend our "
        "detector with the 5 missing signals; direct numeric port). It is "
        "the cheapest CARF upgrade on the list and would specifically catch "
        "the 'fluent but vague' hallucinations our entity/statistic check "
        "misses.",
        italic=True,
    )

    # --------------------------- §8 Provenance ledger / EpistemicState ---------------------------
    h1(doc, "8. EpistemicState ↔ claim_provenance ledger")

    p(
        doc,
        "Plain language: CARF requires every analytical output to carry an "
        "EpistemicState record describing what was used to produce it. "
        "Climatenews has the same idea, implemented as a Postgres ledger "
        "called claim_provenance. Every row records the model, prompt "
        "name/version/fingerprint, retrieval strategy, source article ids, "
        "and confidence. The shape is different but the role is identical.",
    )

    h2(doc, "8.1 The schema (021_claim_provenance.sql)")
    bullet(doc, "Linkage: at least one of {claim_id, url_analysis_id, article_id, deep_search_session_id, cynefin_classification_id} must be set (CHECK constraint).")
    bullet(doc, "extraction_method enum: 'url_analysis_claim_extraction' | 'deep_search_synthesis' | 'cynefin_classification' | 'hallucination_check' | 'article_ingestion_enrichment'.")
    bullet(doc, "Prompt provenance: prompt_name, prompt_version, prompt_fingerprint (CHAR(16)).")
    bullet(doc, "Retrieval: retrieval_strategy (human-readable), source_article_ids (JSONB array).")
    bullet(doc, "Quality: hallucination_score (0–1), confidence (0–1).")
    bullet(doc, "Metadata: raw_metadata JSONB catch-all.")

    h2(doc, "8.2 Gap vs CARF's EpistemicState")
    bullet(doc, "CARF's EpistemicState splits uncertainty into epistemic (model doesn't know enough) and aleatoric (the world is noisy) — climatenews collapses both into a single 'confidence' float.")
    bullet(doc, "CARF covers six analytical layers; climatenews covers five extraction methods. Roughly equivalent but the upstream's layering is more formal.")
    bullet(doc, "CARF can produce a TLA+ formal-verifiable state graph; climatenews relies on Postgres CHECK constraints + integration tests instead.")
    bullet(doc, "Both share the structurally-important detail that provenance writes are best-effort and non-load-bearing — provenance.py:127–136 logs a warning and returns None on failure rather than breaking the calling pipeline. This is right; an audit-trail bug must not crash a user response.")

    # --------------------------- §9 Guardian / editorial gate ---------------------------
    h1(doc, "9. Guardian layer ↔ EditorialGate")

    p(
        doc,
        "Plain language: CARF wraps every output in a Guardian policy "
        "layer that enforces organisational rules before the output reaches "
        "the user. Climatenews has an EditorialGate that does the same job "
        "for one specific surface — publishing an analysed article — but "
        "with Python-hardcoded rules instead of CARF's externalised YAML + "
        "CSL-Core + OPA stack.",
    )

    h2(doc, "9.1 What climatenews does")
    p(
        doc,
        "File: src/backend/app/domains/intelligence/editorial_gate.py. "
        "EditorialGate.evaluate() returns one of three decisions:",
    )
    bullet(doc, "PUBLISH — reliability_score ≥ 60, no disputed claims, high-tier source.")
    bullet(doc, "HOLD — reliability_score 30–59, OR some disputed claims.")
    bullet(doc, "ESCALATE — reliability_score < 30, OR majority disputed, OR low source tier.")
    p(
        doc,
        "Risk factors (free-form list returned to the caller): very-low "
        "reliability, majority disputed, public-tier-with-low-reliability, "
        "low verification confidence, preprint-not-peer-reviewed. The gate "
        "is invoked at the end of url_analyses processing.",
    )

    h2(doc, "9.2 Gap vs CARF Guardian")
    bullet(doc, "No YAML externalisation. CARF Guardian rules live in config/policies.yaml + config/policies/ (35 formal CSL-Core rules). Editorial team cannot edit our rules without code review.")
    bullet(doc, "No CSL-Core formal-rule engine. We use Python if-branches.")
    bullet(doc, "No OPA integration. The audit explicitly declined OPA as overkill, but YAML externalisation (item #7) was recommended — 2 days of work to declare editorial rules as YAML evaluated in editorial_gate.")
    bullet(doc, "No fail-closed determinism guarantee. CARF Guardian is auditable as fail-closed; our gate's behaviour on a missing reliability_score is to fall through to score=0, which is effectively ESCALATE but only by accident.")

    # --------------------------- §10 Drift / Bias / Convergence ---------------------------
    h1(doc, "10. Monitoring stack — drift shipped, bias missing")

    h2(doc, "10.1 Drift detection")
    p(
        doc,
        "CARF tracks KL divergence on the routing distribution (how often "
        "each Cynefin domain fires). climatenews tracks KL divergence on "
        "source mix — does one publisher's share spike unexpectedly, did a "
        "feed go dark.",
    )
    p(
        doc,
        "File: src/backend/app/domains/intelligence/drift_detection.py. "
        "DriftReport carries metric, kl_divergence (nats), verdict, window "
        "sizes, counts, and top_shifts. Verdict thresholds (calibrated on "
        "the 2026-04 → 2026-05 production seed):",
    )
    bullet(doc, "KL < 0.10 — stable; no action.")
    bullet(doc, "0.10 ≤ KL < 0.25 — minor drift; log for trend analysis.")
    bullet(doc, "0.25 ≤ KL < 0.50 — notable drift; flag in ops dashboard.")
    bullet(doc, "KL ≥ 0.50 — significant drift; page on-call.")
    p(
        doc,
        "Smoothing: ε = 1e-6 added to every count before normalising to "
        "avoid divide-by-zero when a previously-zero source appears. "
        "Gap: we only track one DriftReport (source_mix). The 2026-05-18 "
        "audit item #9 recommends adding a second one for claim-type drift "
        "(attribution / projection / baseline) to catch narrative shifts "
        "the source mix cannot see. 1 day of work.",
    )

    h2(doc, "10.2 Bias auditing — missing")
    p(
        doc,
        "CARF runs three chi-squared tests on its memory corpus: domain "
        "distribution (χ² p < 0.05), quality disparity (>20% gap), "
        "approval disparity (>15% gap). H41 benchmark: 100% detection "
        "accuracy.",
    )
    p(
        doc,
        "Climatenews has no equivalent module. The audit (item #2, 2 days) "
        "recommends a chi-square over claim-type × verdict counts to give a "
        "defensible 'we audit ourselves' number for the truth-machine grade. "
        "Climate corpora inherit ideological skew from source mix; without "
        "this auditor, that skew is invisible.",
    )

    h2(doc, "10.3 Convergence (plateau) detection — missing")
    p(
        doc,
        "CARF detects retraining plateaus via the Convergence Monitor (H42, "
        "100% accuracy). Climatenews does not retrain models in-platform, "
        "so the convergence monitor is not directly applicable today — but "
        "it becomes relevant the moment we train our own embeddings or "
        "verifier. Deferred until we own a training loop.",
    )

    # --------------------------- §11 RAG and Three-Layer retrieval ---------------------------
    h1(doc, "11. Three-Layer RAG — the cleanest match")

    p(
        doc,
        "Plain language: the one CARF idea that is fully shipped in "
        "climatenews and serving every search and chat query in "
        "production. Three retrieval channels — vector similarity, "
        "full-text search, entity-graph traversal — each return a ranked "
        "list, and the three lists are fused with Reciprocal Rank Fusion. "
        "No additional CARF feature would change this surface today.",
    )

    h2(doc, "11.1 The implementation")
    p(
        doc,
        "File: src/backend/app/domains/intelligence/hybrid_rag_service.py. "
        "HybridRAGService.retrieve() runs:",
    )
    bullet(doc, "_semantic_search — pgvector HNSW cosine on articles.embedding (OpenAI text-embedding-ada-002, 1536-dim). ef_search=40 for chat, 100 for deep-search.")
    bullet(doc, "_fulltext_search — websearch_to_tsquery against articles.search_tsv (multilingual GIN index).")
    bullet(doc, "_graph_search — BFS from entities matched in the query to articles that mention those entities (JSONB on claims.entities).")
    bullet(doc, "_reciprocal_rank_fusion — 1 / (k + rank) per list, summed per document, top fifteen returned.")

    h2(doc, "11.2 Where the depth differs")
    bullet(doc, "CARF's symbolic layer is a Neo4j forward-chaining engine; ours is a JSONB BFS. CARF can do real graph-pattern queries; we cannot.")
    bullet(doc, "CARF's vector layer uses domain-specific embeddings trained for the climate / governance corpus; we use generic ada-002. Item to revisit: text-embedding-3-large is roughly 2× the cost for ada-002, with materially better recall on climate-domain queries.")
    bullet(doc, "Fusion: both use RRF, so this layer is genuinely the same idea.")

    # --------------------------- §12 EU AI Act, Article 50 ---------------------------
    h1(doc, "12. EU AI Act — proxy partial, Article 50 disclosure shipped")

    p(
        doc,
        "Plain language: the EU AI Act takes effect in waves through 2026 "
        "and 2027. Article 50 (effective 2 Aug 2026) requires AI-generated "
        "content to be disclosed as such. Articles 9 / 12 / 13 / 14 cover "
        "risk management, record-keeping, transparency, and human oversight "
        "for higher-risk systems.",
    )

    h2(doc, "12.1 Article 50 disclosure (shipped)")
    p(
        doc,
        "Every AI-generated artefact on the platform (claim verdict, "
        "deep-search synthesis, plain-language interpretation, country "
        "biome narrative) carries a Schema.org CreativeWork JSON-LD block "
        "in the page HTML with model identifier, prompt fingerprint, and "
        "timestamp. This is the platform's own implementation — not "
        "proxied through CARF.",
    )

    h2(doc, "12.2 Articles 9/12/13/14 (proxied stub)")
    p(
        doc,
        "carf_integration.check_eu_ai_compliance() returns Article 9, 12, "
        "13, 14 verdicts when the CARF service is reachable; otherwise "
        "returns {'compliant': True, 'note': 'CARF compliance check "
        "unavailable'}. The 'compliant: True' default is structurally "
        "optimistic and should be flipped to 'compliant: unknown' for "
        "honesty — a small fix worth doing.",
    )

    # --------------------------- §13 What CARF gives us that we still need ---------------------------
    h1(doc, "13. What CARF would still give us — ranked")

    p(
        doc,
        "The 2026-05-18 audit produced an ordered list of CARF components "
        "worth porting, with effort estimates. Reproducing it here with "
        "current grades:",
    )

    kv_table(
        doc,
        ["#", "Component", "Effort", "Current grade", "Why it matters"],
        [
            ["1", "8-signal hallucination fusion", "1 day", "PARTIAL (3/8)", "Catches 'fluent but vague' hallucinations our 3-signal check misses. Direct numeric port; cheapest CARF upgrade."],
            ["2", "Chi-squared bias auditor", "2 days", "MISSING", "Gives a defensible 'we audit ourselves' number. No equivalent in climatenews."],
            ["3", "Three-mode Bayesian (approximate / PyMC MCMC / cached) with 90% credible intervals + epistemic vs aleatoric", "3 days", "PARTIAL (single point estimate)", "Lifts the honesty axis materially. Today the credibility score is one number; the upgrade returns a (lo, hi) band and splits 'we don't know' from 'the world is noisy'."],
            ["4", "DoWhy refutation battery (placebo, random common cause, 80% subset, E-value)", "4 days", "PARTIAL (LLM-only)", "Highest-leverage upgrade for the causal axis. Even one DoWhy run per headline claim on the temperature/CO₂/IPCC tabular series gives one real causal number per article."],
            ["5", "MAP–PRICE–RESOLVE governance", "5 days", "MISSING", "PRICE = per-article token/$ cost; MAP = which climate sub-domains an article touches; RESOLVE = editorial-policy-conflict surfacing. Strong differentiator for truth-machine framing."],
            ["6", "Plotly visualisations (CausalDAG, BayesianPanel, CynefinMatrix, …)", "4 days", "MISSING", "Visually justifies the scoring instead of just printing a number. Three frontend additions, mountable on /analyze and /articles/[slug]."],
            ["7", "CSL policy YAML scaffold (editorial rules → YAML)", "2 days", "PARTIAL (Python-hardcoded)", "Editorial team can edit policies without code review. Matches the upstream Guardian pattern."],
            ["8", "Versioned prompts.yaml — self-correction prompt", "1 day", "PARTIAL (registry exists, self-correction missing)", "Hardens deep-search and analyse pipelines with a 'process rejection → retry → escalate after N' reflection step."],
            ["9", "Drift-detector multi-dimension (claim-type drift)", "1 day", "PARTIAL (source-mix only)", "Catches narrative shifts (sudden spike in projection-style claims = misinformation campaign) that source-mix drift cannot."],
        ],
    )

    p(
        doc,
        "Total: 23 days of focused work would close the gap on every "
        "component the audit recommends. The cheapest three (items 1, 2, 8) "
        "are 4 days and lift the truth-machine grade visibly.",
    )

    # --------------------------- §14 The pieces we declined ---------------------------
    h1(doc, "14. What we explicitly did not port (and why)")

    kv_table(
        doc,
        ["Upstream piece", "Reason for declining"],
        [
            ["LangGraph StateGraph workflow", "Climatenews already has its own orchestration via services.py + Celery + Cloud Scheduler. LangGraph would be a rewrite, not an improvement."],
            ["OPA / Open Policy Agent", "Overkill for a public fact-check site. YAML policies (item #7 above) would be sufficient if we add policy externalisation."],
            ["HumanLayer approval workflows", "There is no editorial human-in-the-loop UI yet. Deferred until we build a moderation queue."],
            ["Kafka audit log", "Current Postgres audit_events table is sufficient at our scale (~500 provenance rows / day)."],
            ["TLA+ formal specs (StateGraph, EscalationProtocol)", "Not justifiable for the climate-news truth-machine grade. The cost of formal verification at our scope outweighs the benefit; provenance + integration tests are enough."],
            ["17 demo scenarios / 43-hypothesis benchmark harness", "Upstream scenarios cover Scope-3, supply-chain, and manufacturing — none climate-news-specific. We will build our own DeepEval suite (already planned in the truth-machine roadmap)."],
            ["MCP server with 18 cognitive tools", "We expose REST + Next.js; there are no MCP consumers in the climatenews roadmap."],
            ["Firebase Auth", "Climatenews already runs on GCP Cloud SQL with custom JWT auth + Stripe billing. No migration value."],
            ["ChimeraOracle fast-path (pre-trained CausalForestDML, <100ms)", "Defer until the causal pipeline (item #4 above) is real. Pre-training before the LLM-only causal extraction is replaced would be premature optimisation."],
        ],
    )

    # --------------------------- §15 License + provenance ---------------------------
    h1(doc, "15. License and provenance of the port")

    p(
        doc,
        "Upstream license: Business Source License 1.1, converts to Apache "
        "2.0 on 2030-02-19. Production use of upstream source code requires "
        "a commercial license from licensing@cisuregen.com.",
    )
    p(doc, "Climatenews's stance, as recorded in the 2026-05-18 audit and reaffirmed in this report:")
    bullet(doc, "Treat CARF components as algorithmic inspiration, re-implement clean. No upstream source files are vendored into climatenews.")
    bullet(doc, "When a prompt or numeric weight is ported directly (e.g. the cynefin_classifier LLM prompt), the docstring acknowledges the upstream source: 'Prompt adapted from projectcarfcynepic/config/prompts.yaml#router (eljaplacido/projectcarfcynepic — BSL 1.1, both repos same owner so self-permission is fine; track in docs/licenses if the port grows)' — see cynefin_router.py:236–240.")
    bullet(doc, "Trademarks 'CARF' and 'CYNEPIC' belong to Cisuregen. Internal docs refer to them as upstream concepts, not climatenews brand assets.")
    bullet(doc, "The same person owns both repositories, so the practical license risk is low. The discipline is still 'reimplement clean' to keep the artefact distinct.")

    # --------------------------- §16 Honest verdict ---------------------------
    h1(doc, "16. Honest verdict — does climatenews implement CARF?")

    p(
        doc,
        "Yes, structurally. The five CARF ideas that matter most for a "
        "climate-news truth-machine are all present in some form in the "
        "running code:",
    )
    bullet(doc, "Cynefin routing as the entry-point classification (cynefin_router.py).")
    bullet(doc, "Provenance ledger as the substrate of every analytical output (claim_provenance + provenance.py).")
    bullet(doc, "Three-layer RAG with RRF as the retrieval pattern (hybrid_rag_service.py).")
    bullet(doc, "Hallucination sentinel on every AI-generated artefact (hallucination_detector.py).")
    bullet(doc, "EditorialGate as the local Guardian (editorial_gate.py).")

    p(
        doc,
        "No, in depth. Each of the five is shipped as a thinner version of "
        "its upstream counterpart:",
    )
    bullet(doc, "Cynefin labels but does not route into different engines.")
    bullet(doc, "Provenance ledger does not split epistemic and aleatoric uncertainty.")
    bullet(doc, "RAG is identical in concept; vector model is generic ada-002, graph layer is JSONB BFS not Neo4j.")
    bullet(doc, "Hallucination uses 3 signals of 8.")
    bullet(doc, "Guardian rules are Python, not YAML+CSL+OPA.")

    p(
        doc,
        "The four CARF pieces that genuinely do not exist in climatenews "
        "today — chi-squared bias auditor, MAP–PRICE–RESOLVE governance, "
        "DoWhy refutation battery, plotly visualisations — are the ones "
        "that would most visibly close the depth gap. The audit's ranked "
        "list in §13 is the work order.",
    )

    p(
        doc,
        "Net: climatenews is a faithful but reduced application of the "
        "CARF / CYNEPIC architecture, narrowed from a general epistemic-"
        "reasoning research framework to a climate-news verification "
        "platform. The architectural shape is the same; the depth on each "
        "axis is roughly 50–70% of the upstream's. Closing that depth gap "
        "is a focused 20–25 day project against the audit's ranked list "
        "and would convert the platform from 'inspired by CARF' to "
        "'CARF-grade on the climate-news axis'.",
    )

    # --------------------------- Footer ---------------------------
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run(
        "End of report — Climatefacts.ai · 2026-05-25 · "
        "Upstream: github.com/eljaplacido/projectcarfcynepic (BSL 1.1) · "
        "Companion: docs/audits/2026-05-18-projectcarfcynepic-comparison.md"
    )
    r.italic = True
    r.font.size = Pt(9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build()
    print(f"Wrote {out}")

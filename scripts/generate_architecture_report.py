#!/usr/bin/env python
"""Generate Climatefacts.ai architecture + feature report as .docx.

Run:
    python scripts/generate_architecture_report.py

Output:
    docs/reports/Climatefacts-Architecture-Report-2026-05-24.docx
"""
from __future__ import annotations

from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = (
    REPO_ROOT / "docs" / "reports"
    / f"Climatefacts-Architecture-Report-{date.today().isoformat()}.docx"
)


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

TEAL = RGBColor(0x0F, 0x76, 0x6E)
SLATE = RGBColor(0x33, 0x3F, 0x4D)
MUTED = RGBColor(0x6B, 0x72, 0x80)
AMBER = RGBColor(0xB4, 0x53, 0x09)
RED = RGBColor(0xB9, 0x1C, 0x1C)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = TEAL if level <= 2 else SLATE


def add_para(doc: Document, text: str, *, italic: bool = False, bold: bool = False,
             color: RGBColor | None = None, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_bullet(doc: Document, text: str, *, bold_lead: str | None = None,
               level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    if bold_lead:
        lead = p.add_run(bold_lead)
        lead.bold = True
        rest = p.add_run(text)
    else:
        p.add_run(text)


def add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    k = p.add_run(f"{key}: ")
    k.bold = True
    p.add_run(value)


def add_callout(doc: Document, text: str, *, color: RGBColor = TEAL) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = color
    run.font.size = Pt(10)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = SLATE


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------


def build() -> None:
    doc = Document()

    # ---- Title page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Climatefacts.ai")
    title_run.bold = True
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = TEAL

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Technical Architecture & Feature Report")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = SLATE

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Production deploy snapshot — {date.today().isoformat()}\n"
        "GCP project climatenews-495412 (europe-west4) · Owner: CISU Regen"
    )
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = MUTED
    meta_run.italic = True

    add_page_break(doc)

    # ---- 1. Executive Summary ----
    add_heading(doc, "1. Executive Summary", 1)
    add_para(doc,
        "Climatefacts.ai is a climate-news fact-checking + intelligence platform "
        "built around three structural properties competitors cannot match: deterministic "
        "claim verification (no LLM in the verdict path), seven first-class personas "
        "served through the same data with audience-appropriate framings, and a "
        "fully-introspectable agentic chat layer whose action set is single-sourced "
        "and cross-verified by automated tests.")
    add_para(doc,
        "The platform combines a multi-LLM verification pipeline (cross-source "
        "corroboration + numeric grounding) with public regulatory datasets (CDP, "
        "SBTi, Net Zero Tracker, IPCC AR6) and 200+ RSS feeds covering UN-193 "
        "countries. It is positioned for ECGT enforcement (27 Sept 2026) and "
        "CSRD wave-1 reporting cycles, with a Business view mode that converts "
        "scientific framings to fiduciary risk framings for board-facing use.")
    add_callout(doc,
        "Single platform. Seven personas. One source of truth per claim — every "
        "verdict traceable to the disclosure ledger or the news corpus it came from.")

    add_heading(doc, "Key Numbers (production, 2026-05-24)", 2)
    add_bullet(doc, "190+ countries covered (UN-193 minus 3 territories)", bold_lead="Coverage: ")
    add_bullet(doc, "200+ RSS feeds across global, regional, scientific tiers", bold_lead="Feeds: ")
    add_bullet(doc, "10,000+ companies tracked with SBTi/CDP/NZT data", bold_lead="Corporate: ")
    add_bullet(doc, "180 country × scenario × horizon projection rows (IPCC AR6)", bold_lead="Projections: ")
    add_bullet(doc, "11 single-sourced agentic skills, dispatcher-verified", bold_lead="Agentic: ")
    add_bullet(doc, "153 backend pytest + 350 frontend Vitest + Playwright E2E", bold_lead="Tests: ")
    add_bullet(doc, "TypeScript strict, 0 errors", bold_lead="Type safety: ")

    add_page_break(doc)

    # ---- 2. Solution Vision ----
    add_heading(doc, "2. Solution Vision", 1)

    add_heading(doc, "Mission", 2)
    add_para(doc,
        "Make climate information defensible. Every fact, claim, projection, and "
        "corporate disclosure on the platform is traceable to its source. No LLM "
        "is allowed in the verdict path. Users always see how the platform decided "
        "what to show them.")

    add_heading(doc, "Three structural differentiators vs every competitor", 2)
    add_bullet(doc,
        "Single platform serves consumers, journalists, ESG officers, scientists, "
        "policymakers, financial analysts, and business decision-makers — each "
        "with audience-appropriate framings on the same underlying data. "
        "(OWID is science-only. Climate Watch is policy-only. Bloomberg ESG is "
        "finance-only. Nobody else covers all seven.)",
        bold_lead="Persona breadth — ")
    add_bullet(doc,
        "Claim verdicts come from deterministic rules (numeric grounding, "
        "SBTi-validation lookup, ECGT keyword matchers), not from an LLM "
        "summarising its own confidence. LLMs only extract candidate claims; "
        "they never decide truth. Every verdict is reproducible from inputs.",
        bold_lead="LLM-free verdict path — ")
    add_bullet(doc,
        "The 11 agentic skills the chat can emit are defined in ONE Python "
        "registry (SKILLS_REGISTRY). The frontend dispatcher, prompt template, "
        "and pin tests all read from it. Drift between the LLM and the runtime "
        "is impossible — tests block it at CI time.",
        bold_lead="Single-sourced agentic protocol — ")

    add_heading(doc, "Competitive positioning", 2)
    add_table(
        doc,
        ["Capability", "Climatefacts.ai", "OWID", "Climate Watch", "Climate TRACE", "Probable Futures"],
        [
            ["Country profiles", "Yes (6 tabs, dual view)", "Limited", "Yes", "No", "Yes"],
            ["News + claims integration", "Yes (200+ RSS)", "No", "No", "No", "No"],
            ["Corporate verification", "Yes (CDP/SBTi/NZT)", "No", "Partial", "No", "No"],
            ["Multi-LLM cross-verification", "Yes", "N/A", "N/A", "N/A", "N/A"],
            ["IPCC AR6 scenarios per country", "Yes (SSP1-2.6/2-4.5/3-7.0)", "Yes", "No", "No", "Partial"],
            ["AOI threshold alerts", "Yes", "No", "No", "No", "No"],
            ["Business view mode", "Yes (CSRD/IFRS S2 chips)", "No", "No", "No", "No"],
            ["Agentic chat with provenance", "Yes (11 skills)", "No", "No", "No", "No"],
            ["Embed widgets", "Yes", "Yes", "Yes", "No", "Yes"],
            ["ECGT-aware claim flagging", "Yes", "No", "No", "No", "No"],
        ],
    )

    add_page_break(doc)

    # ---- 3. User Personas ----
    add_heading(doc, "3. User Personas", 1)
    add_para(doc,
        "Seven personas are served. The same data feeds each, but the framing, "
        "default view, and quota tier differ. Every page that surfaces a number "
        "should pass the test: 'would the right framing change if this user were "
        "a board chair vs a sixth-grader vs a climate scientist?' If yes, the "
        "page should respect that via the view-mode toggle or persona-appropriate "
        "defaults.")

    personas = [
        ("3.1 Consumer", "Citizens checking climate news + understanding what "
                         "their country's climate position actually is.", "Free / Basic",
         ["Country Climate Passport (Public view)",
          "Plain-language interpretation sentences",
          "URL analyser for claims they encounter",
          "Article credibility chips",
          "Saved articles (3 free, unlimited Basic+)"]),
        ("3.2 Journalist", "Reporters needing evidence-backed claims, source diversity, "
                          "deep search across the corpus + climate projections "
                          "for their region.", "Basic / Professional",
         ["Deep Search (search + compare modes)",
          "Claim Ledger per country",
          "URL analyser with structured failure surface",
          "Country comparisons by region/topic",
          "Embed widgets for online articles",
          "Methodology page (defensibility for editorial review)"]),
        ("3.3 ESG / Sustainability Officer", "Sustainability teams gathering corporate "
                                             "evidence, validating supplier claims, preparing "
                                             "CSRD/IFRS S2 disclosures.", "Professional / Enterprise",
         ["Corporate Climate Tracker (/companies)",
          "Per-company disclosure trail (CDP / SBTi / NZT)",
          "Business view: compliance chips (CSRD / IFRS S2 / TCFD)",
          "Claim verification analyzer (ECGT Article 4)",
          "Country Passport in Business view for supplier-country risk",
          "Bookmark unlimited corporate claims for audit",
          "PDF/CSV export"]),
        ("3.4 Climate Scientist / Researcher", "Researchers needing methodology transparency, "
                                                "raw indicators, and projection scenarios.", "Basic",
         ["Methodology page (prompt registry, fingerprints, calibration)",
          "Country indicators (raw)",
          "IPCC AR6 projections per scenario",
          "Source-tier raw scores",
          "Claim provenance with confidence bands"]),
        ("3.5 Policymaker", "Policy advisors evaluating national positions, regulatory "
                            "compliance, ambition gaps.", "Professional",
         ["Country Passport overview tab",
          "Region comparisons + UNFCCC NDC indicators",
          "Climate Action Tracker integration",
          "Multi-country deep-search compare",
          "AOI alerts for policy-relevant thresholds (renewable share, anomaly)"]),
        ("3.6 Financial Analyst", "Equity research analysts assessing climate physical-risk "
                                  "exposure, regulatory headwinds, transition risk.", "Professional / Enterprise",
         ["Country Passport in Business view (physical-risk framing)",
          "IPCC AR6 SSP projections (scenario analysis)",
          "Corporate Climate Tracker (transition-risk audit)",
          "Region/sector deep-search compare",
          "Embed widgets for research reports",
          "Export CSV/PDF"]),
        ("3.7 Business Decision-Maker", "C-suite, board, procurement, Chief Risk Officers "
                                         "vetting climate-related capex, supplier claims, "
                                         "market expansion.", "Enterprise",
         ["Business view toggle on Country Passport + Company Detail",
          "Compliance-framework chips (CSRD / IFRS S2 / TCFD / TNFD)",
          "ECGT Article 4 risk warnings on offset-based claims",
          "Risk-framed (not science-framed) plain-language sentences",
          "Board-ready exports",
          "URL persistent ?view=business for sharing boardroom links"]),
    ]
    for title, who, tier, features in personas:
        add_heading(doc, title, 2)
        add_kv(doc, "Who", who)
        add_kv(doc, "Default tier", tier)
        add_kv(doc, "Primary features", "")
        for f in features:
            add_bullet(doc, f)

    add_page_break(doc)

    # ---- 4. Feature Catalog ----
    add_heading(doc, "4. Feature Catalog", 1)
    add_para(doc,
        "Every public-facing surface, what it does, which persona it primarily "
        "serves, and how it works under the hood. Internal flags + service-only "
        "endpoints are documented in the operator runbook.")

    features = [
        ("4.1 World Climate Map (/map)",
         "Interactive Leaflet map with 4 layers (article density, temperature "
         "anomaly, climate risk, source diversity). Click a country → "
         "right-panel country card → 'Open passport' navigates to Country Passport.",
         "Consumer, Journalist, Policymaker",
         "GET /api/map/country-stats, GET /api/map/topic-density, "
         "GET /api/map/source-coverage."),
        ("4.2 Country Climate Passport (/country/[code])",
         "6 tabs: Overview / News / Climate Data / Projections / Sources / Claim "
         "Ledger. Each tab is URL-deeplinkable. Public ↔ Business view toggle "
         "swaps plain-language framings and stamps compliance chips per KPI.",
         "All personas",
         "GET /api/map/country/{cc}/detail, GET /api/map/country/{cc}/climate-data, "
         "GET /api/map/country/{cc}/projections, GET /api/map/country/{cc}/claim-ledger."),
        ("4.3 Corporate Climate Tracker (/companies)",
         "Listing of 10,000+ companies with SBTi/CDP/NZT data. Filter + search. "
         "Per-company page shows disclosures, scope 1/2/3 emissions, "
         "verified claims, ECGT/CSRD compliance chips in Business view, and a "
         "claim verification analyzer.",
         "ESG officer, Business decision-maker, Financial analyst",
         "GET /api/companies, GET /api/companies/{ticker}, "
         "POST /api/companies/{ticker}/analyze (rate-limited, "
         "ECGT-aware), GET/POST /api/companies/admin/sync/{source} "
         "(token-gated background sync)."),
        ("4.4 Deep Search (/deep-search)",
         "Two modes: Search (evidence-grounded answer + citations) and Compare "
         "(two queries side-by-side). Refinement chips, clarification prompts, "
         "compare-charts, methodology drawer. URL state via useUrlState.",
         "Journalist, Researcher, Policymaker",
         "POST /api/deep-search, POST /api/deep-search/compare."),
        ("4.5 Article Search (/search)",
         "Tag, country, credibility, category filters. Multilingual FTS. "
         "URL-persistent filter set.",
         "All",
         "GET /api/articles with structured filters."),
        ("4.6 URL Analysis (/analyze)",
         "Submit a URL for fact-checking. Pipeline: content extraction → claim "
         "extraction (multi-LLM) → numeric grounding → source credibility → "
         "calibration label. 15 structured failure reasons with persona-appropriate "
         "recovery surfaces.",
         "Consumer, Journalist",
         "POST /api/analyze-url, GET /api/analyze-url/{id}."),
        ("4.7 Methodology Page (/methodology)",
         "Read-only transparency surface. Sections: prompts inventory, "
         "sustainability formula, calibration metrics, hallucination rates, "
         "drift verdicts, source tiers, corporate-verification taxonomy.",
         "Auditor, Journalist, Researcher",
         "GET /api/methodology (bundled), plus separate calibration / drift / "
         "indicator endpoints."),
        ("4.8 Agentic Chat",
         "Site-wide chat panel. The LLM emits 0-3 'action' suggestions per "
         "answer, drawn from a registry of 11 typed skills. Auto skills navigate "
         "immediately; confirm skills (4) go through a confirmation modal "
         "before consuming quota or mutating state.",
         "All",
         "POST /api/chat, dispatcher in chatActionDispatcher.ts; "
         "skills served via GET /api/skills."),
        ("4.9 AOI Alerts (/dashboard/aoi)",
         "Subscribe to threshold-crossing alerts per country + variable "
         "(temperature anomaly, renewable share, CO₂ per capita). "
         "Daily Cloud Scheduler polls indicators; debounce prevents email-bombing "
         "on steady-state crossings.",
         "Journalist, ESG officer, Policymaker",
         "POST /api/aoi-subscriptions, DELETE /api/aoi-subscriptions/{id}, "
         "POST /api/scheduler/aoi-poll (cron-only)."),
        ("4.10 Embed Widgets (/embed/*)",
         "iframe-embeddable widgets (Country Passport mini, claim card, "
         "deep-search result). For journalists embedding in online articles "
         "and analysts in research reports.",
         "Journalist, Financial Analyst",
         "Pure SSR, sized for 600×400 default."),
    ]
    for title, what, who, how in features:
        add_heading(doc, title, 2)
        add_kv(doc, "What", what)
        add_kv(doc, "Serves", who)
        add_kv(doc, "Backend", how)

    add_page_break(doc)

    # ---- 5. Architectural Components ----
    add_heading(doc, "5. Architectural Components", 1)
    add_para(doc, "The platform is split into Content, Intelligence, and Trust "
                  "domains. Each domain has its own module hierarchy, tests, and "
                  "API surface. The Trust domain wraps Content + Intelligence with "
                  "the provenance / calibration / verification layers that make "
                  "every claim defensible.")

    add_heading(doc, "5.1 Frontend (Next.js 14)", 2)
    add_bullet(doc, "Next.js 14 App Router with React 18", bold_lead="Framework: ")
    add_bullet(doc, "Tailwind CSS for design system; lucide-react for icons", bold_lead="Styling: ")
    add_bullet(doc, "Custom useUrlState hook for URL-persistent filter state across pages", bold_lead="State: ")
    add_bullet(doc, "Suspense boundaries around all useSearchParams callers (Phase 8 fix)", bold_lead="Rendering: ")
    add_bullet(doc, "plainLanguage.ts — shared formatter library with public + business framings", bold_lead="Layer: ")
    add_bullet(doc, "Vitest (350 tests) + Playwright E2E (60+ scenarios)", bold_lead="Tests: ")
    add_bullet(doc, "TypeScript strict mode, 0 errors", bold_lead="Types: ")

    add_heading(doc, "5.2 Backend (FastAPI on Cloud Run)", 2)
    add_bullet(doc, "FastAPI 0.109+ on Python 3.11", bold_lead="Framework: ")
    add_bullet(doc, "Domain-driven design: Content / Intelligence / Trust", bold_lead="Architecture: ")
    add_bullet(doc, "Async/await throughout; BackgroundTasks for long-running adapter syncs", bold_lead="Concurrency: ")
    add_bullet(doc, "Pydantic v2 for request/response models", bold_lead="Validation: ")
    add_bullet(doc, "OpenTelemetry tracing → Cloud Logging", bold_lead="Observability: ")
    add_bullet(doc, "1800s request timeout (bumped from 300s for adapter syncs)", bold_lead="Cloud Run: ")
    add_bullet(doc, "min-instances=1 (warm) so background tasks complete", bold_lead="Scaling: ")

    add_heading(doc, "5.3 Database (PostgreSQL 15 on Cloud SQL)", 2)
    add_bullet(doc, "Cloud SQL instance climatenews-postgres (europe-west4)", bold_lead="Host: ")
    add_bullet(doc, "pgvector extension for HNSW article embedding index", bold_lead="Extensions: ")
    add_bullet(doc, "Multilingual FTS (German / French / English / Finnish)", bold_lead="Search: ")
    add_bullet(doc, "36 versioned SQL migrations, idempotent runner via cloud-sql-proxy", bold_lead="Schema: ")
    add_bullet(doc, "schema_migrations_applied tracker (Phase 8 addition)", bold_lead="Tracking: ")

    add_heading(doc, "5.4 Caching & Rate Limiting", 2)
    add_bullet(doc, "Redis backing IP-based + tier-based daily/monthly counters", bold_lead="Cache: ")
    add_bullet(doc, "Tier-aware throttles on URL analysis + claim verification + chat", bold_lead="Limits: ")
    add_bullet(doc, "Structured 429 envelope with quota envelope shape (tier/used/limit/upgrade_url)", bold_lead="Errors: ")

    add_heading(doc, "5.5 Async Processing", 2)
    add_bullet(doc, "Celery + Redis broker for verification, URL analysis, content discovery", bold_lead="Queue: ")
    add_bullet(doc, "FastAPI BackgroundTasks for short-running tasks (adapter sync)", bold_lead="Inline: ")
    add_bullet(doc, "Cloud Scheduler triggers daily/hourly cron jobs via authenticated HTTP", bold_lead="Cron: ")

    add_heading(doc, "5.6 AI / LLM Layer", 2)
    add_bullet(doc, "Multi-LLM verifier — DeepSeek (primary) + Anthropic Claude (secondary)", bold_lead="Cross-source: ")
    add_bullet(doc, "Numeric grounding check — pure-function defence against shared hallucination", bold_lead="Grounding: ")
    add_bullet(doc, "Cynefin router classifies query into clear / complicated / complex / chaotic", bold_lead="Routing: ")
    add_bullet(doc, "Prompt registry with SHA-256 fingerprints surfaced on /methodology", bold_lead="Prompts: ")
    add_bullet(doc, "Pure Python verdict path — LLMs never decide truth, only extract candidates", bold_lead="Verdict: ")

    add_heading(doc, "5.7 Data Ingestion", 2)
    add_bullet(doc, "rss_feed_registry — 200+ feeds across UN-193 countries via Google News + native feeds", bold_lead="Feeds: ")
    add_bullet(doc, "SBTi public Google Sheets adapter (2026 schema)", bold_lead="SBTi: ")
    add_bullet(doc, "CDP / Net Zero Tracker adapters with graceful 'data source moved' warnings", bold_lead="Corporate: ")
    add_bullet(doc, "IPCC AR6 SSP1-2.6/2-4.5/3-7.0 seed for 20 countries × 3 horizons", bold_lead="Projections: ")
    add_bullet(doc, "Country indicators from World Bank, IRENA, OWID, UNFCCC NDC, ND-GAIN, Climate TRACE", bold_lead="Indicators: ")
    add_bullet(doc, "Source profile registry with credibility tiers (Tier 1 / 2 / 3)", bold_lead="Sources: ")

    add_heading(doc, "5.8 Authentication", 2)
    add_bullet(doc, "JWT access + rotating refresh tokens (stateful sessions table)", bold_lead="Tokens: ")
    add_bullet(doc, "Google OAuth", bold_lead="SSO: ")
    add_bullet(doc, "Stripe billing integration for tier upgrades", bold_lead="Billing: ")
    add_bullet(doc, "Replay-detection on refresh token rotation (cascade-revoke)", bold_lead="Security: ")

    add_heading(doc, "5.9 Observability", 2)
    add_bullet(doc, "OpenTelemetry — trace_id + request_id surfaced in 5xx response bodies", bold_lead="Tracing: ")
    add_bullet(doc, "JSON-structured logs via structlog → Cloud Logging", bold_lead="Logs: ")
    add_bullet(doc, "Performance audit pipeline produces dated audits under docs/audits/", bold_lead="Audits: ")
    add_bullet(doc, "Drift detection on claim categorisation distributions", bold_lead="Drift: ")

    add_heading(doc, "5.10 Deployment (GCP)", 2)
    add_bullet(doc, "Project climatenews-495412, region europe-west4", bold_lead="Project: ")
    add_bullet(doc, "Cloud Run services: climatenews-api + climatenews-frontend", bold_lead="Runtime: ")
    add_bullet(doc, "Cloud Build pipeline runs migrations as step 0 of every deploy", bold_lead="CI/CD: ")
    add_bullet(doc, "Cloud SQL Auth Proxy for migration step (cloud-sql-proxy v2.11.4)", bold_lead="DB access: ")
    add_bullet(doc, "Cloud Scheduler 7 cron jobs (ingest, RSS, verify, retry, feeds, translate, AOI)", bold_lead="Schedule: ")
    add_bullet(doc, "Secret Manager: 15 secrets (DB, JWT, Anthropic, OpenAI, Stripe, OAuth, sync token)", bold_lead="Secrets: ")
    add_bullet(doc, "GitHub trigger fires on push to main", bold_lead="Trigger: ")

    add_page_break(doc)

    # ---- 6. Trust & Transparency Layer ----
    add_heading(doc, "6. Trust & Transparency Layer", 1)
    add_para(doc, "The single most important architectural commitment: every "
                  "claim, verdict, and projection is traceable to its source. The "
                  "Trust domain wraps the Content + Intelligence domains with the "
                  "components below.")

    add_heading(doc, "6.1 Claim Provenance", 2)
    add_para(doc, "Every extracted claim is stored with its source article ID, "
                  "extraction prompt fingerprint, extraction confidence, multi-LLM "
                  "corroboration score, and a calibration label when verified. "
                  "/methodology surfaces the aggregated metrics live.")

    add_heading(doc, "6.2 Multi-LLM Cross-Verification", 2)
    add_para(doc, "verify_claims() runs two extractors in parallel and computes "
                  "Jaccard similarity on the normalised claim text. Claims covered "
                  "by both extractors get full confidence; primary-only claims get "
                  "confidence × 0.7 (configurable penalty). agreement_score is a "
                  "per-document calibration signal.")

    add_heading(doc, "6.3 Numeric Grounding Check (Phase 8 MVP)", 2)
    add_para(doc, "extract_numbers() + check_numeric_grounding() — pure functions "
                  "that tokenise numeric values + units (°C, ppm, %, GtCO2e, TWh, "
                  "USD-bn, mm, etc.) from a claim and verify each appears in the "
                  "evidence within 1% relative tolerance. Catches the failure "
                  "mode where both LLMs hallucinate the same wrong number.")

    add_heading(doc, "6.4 Source Credibility Tiers", 2)
    add_para(doc, "Sources are classified Tier 1 (scientific: IPCC, Nature, "
                  "scientific journals), Tier 2 (research-grade: Carbon Brief, "
                  "Inside Climate News), Tier 3 (public news: Reuters, NYT, "
                  "Guardian). Article credibility is a weighted average of the "
                  "source tier and the per-article verification results.")

    add_heading(doc, "6.5 Calibration Labels", 2)
    add_para(doc, "Auditors and registered users can label URL analyses with "
                  "their independent verdict. The labelled corpus produces a "
                  "calibration curve published on /methodology (binned by "
                  "confidence band, plotting actual accuracy).")

    add_heading(doc, "6.6 Methodology Snapshot", 2)
    add_para(doc, "Single bundled endpoint GET /api/methodology returns prompts "
                  "(with SHA-256 fingerprints + rationale), sustainability formula "
                  "components + confidence bands, indicator coverage, calibration "
                  "labels, hallucination rates, drift verdicts. Auditors can pin "
                  "a snapshot to a date+commit for record-keeping.")

    add_heading(doc, "6.7 AI Provenance Disclosure", 2)
    add_para(doc, "Every AI-generated artefact (claim verdict, deep-search "
                  "synthesis, plain-language interpretation) carries a "
                  "Schema.org CreativeWork JSON-LD block in the page HTML with "
                  "model identifier, prompt fingerprint, and timestamp. Aligns "
                  "with EU AI Act Article 50 (effective 2 Aug 2026) requirements.")

    add_page_break(doc)

    # ---- 7. Agentic Skills Architecture ----
    add_heading(doc, "7. Agentic Skills Architecture", 1)
    add_para(doc, "The agentic chat panel can suggest 0-3 actions per answer. "
                  "Every action is a typed Skill drawn from a single Python "
                  "registry. The architecture exists because the failure mode of "
                  "drift between 'what the LLM thinks it can do' and 'what the "
                  "frontend can actually execute' is silent — the LLM emits an "
                  "action, the dispatcher ignores it, and the user thinks the "
                  "chat is broken. We block that at CI time.")

    add_heading(doc, "7.1 Skill Registry (Single Source of Truth)", 2)
    add_para(doc, "src/backend/app/domains/intelligence/skills.py — SKILLS_REGISTRY "
                  "dict, 11 entries, each a frozen Skill dataclass with name, "
                  "description, mode (auto | confirm), parameters tuple, and "
                  "target_surfaces tuple. Adding a new skill = one entry here "
                  "+ matching dispatcher handler + bump test pin.")

    add_heading(doc, "7.2 Auto vs Confirm Modes", 2)
    add_para(doc, "Auto skills (7) navigate immediately on click — they're pure "
                  "URL transitions the user can always undo with the browser back "
                  "button. Confirm skills (4) consume quota or mutate server "
                  "state; the dispatcher MUST run them through a confirmation modal. "
                  "Dispatcher fails closed: if a host doesn't wire a confirmation "
                  "callback, confirm skills decline silently rather than execute.")

    add_heading(doc, "7.3 Cross-Source Pin Tests", 2)
    add_para(doc, "test_agentic_skill_pin.py reads SKILLS_REGISTRY, the chat "
                  "prompt template, and chatActionDispatcher.ts's ChatActionType "
                  "union. Asserts all three are perfectly aligned. A drift "
                  "between any two fails CI and blocks merge.")

    add_heading(doc, "7.4 The 11 Skills", 2)
    skills = [
        ("navigate", "auto", "Open a platform route", "{path}"),
        ("analyze_url", "confirm", "Submit a URL for fact-checking analysis", "{url}"),
        ("apply_search_filters", "auto", "Filter /search and navigate there", "{q, credibility, country, tags, category}"),
        ("apply_map_filters", "auto", "Filter the world map", "{country, layer}"),
        ("open_methodology_section", "auto", "Jump to a /methodology section", "{section}"),
        ("open_country", "auto", "Open a country's climate panel on the map", "{code}"),
        ("start_deep_search", "auto", "Launch deep research", "{q}"),
        ("bookmark_article", "confirm", "Save an article (consumes saved-articles quota)", "{article_id}"),
        ("start_calibration_label", "confirm", "Submit a calibration rating", "{url_analysis_id}"),
        ("open_company", "auto", "Open a company's climate disclosure profile", "{ticker}"),
        ("verify_corporate_claim", "confirm", "Verify a corporate climate claim (ECGT/SBTi)", "{ticker, claim_text}"),
    ]
    add_table(doc, ["Skill", "Mode", "Description", "Parameters"], skills)

    add_page_break(doc)

    # ---- 8. Compliance & Regulation ----
    add_heading(doc, "8. Compliance & Regulation", 1)
    add_para(doc, "The platform's data model + verification rules are designed "
                  "around the specific regulatory regimes that take effect in "
                  "2026-2028. Each regime maps to specific feature surfaces.")

    regs = [
        ("8.1 ECGT (Empowering Consumers for Green Transition Directive)",
         "Effective 27 Sept 2026 across EU.",
         "Prohibits offset-based 'climate neutral' product claims without "
         "independent verification. Climatefacts.ai flags any offset-marker "
         "claim deterministically; Business view surfaces an explicit "
         "ECGT Article 4 warning on every offset_based_claims disclosure."),
        ("8.2 CSRD (Corporate Sustainability Reporting Directive)",
         "Wave 1: 2025 fiscal year, reporting in 2026 (ongoing).",
         "Companies' Scope 1+2+3 emissions disclosures aligned with CDP "
         "structure. Business view stamps a CSRD chip on every disclosure "
         "with Scope 1+2 data and on country physical-risk indicators."),
        ("8.3 IFRS S2",
         "Jurisdictional adoption ongoing (Chile, Qatar, Mexico mandatory 2026).",
         "Climate-related disclosures + scenario analysis. Country Passport "
         "Business view stamps IFRS S2 chip on temperature anomaly + "
         "climate risk score; ProjectionsPanel directly enables S2 "
         "scenario-analysis (1.5°C / 2°C / 3°C bracket)."),
        ("8.4 TCFD",
         "Recommended framework, mandatory in some jurisdictions.",
         "Physical-risk + transition-risk disclosure. TCFD chip stamps on "
         "verified scope assurance + climate-risk score."),
        ("8.5 TNFD",
         "Voluntary as of 2026, mandatory pathway 2027+.",
         "Nature-related financial disclosure. TNFD chip stamps on "
         "biodiversity-adjacent indicators in country indicator set."),
        ("8.6 SBTi Corporate Net-Zero Standard",
         "Industry standard, not legally mandatory but widely required "
         "by financiers + procurement.",
         "Net-zero claim verification routes through SBTi validation status. "
         "Companies not on the SBTi registry get 'disputed' verdict on net-zero "
         "claims; the evidence URL points back to the public SBTi registry."),
        ("8.7 EU AI Act Article 50",
         "Effective 2 Aug 2026.",
         "AI provenance disclosure on all generated content. Schema.org "
         "CreativeWork JSON-LD on every page surfacing AI-generated text."),
    ]
    for title, when, what in regs:
        add_heading(doc, title, 2)
        add_kv(doc, "When", when)
        add_kv(doc, "How the platform serves it", what)

    add_page_break(doc)

    # ---- 9. Tier Design ----
    add_heading(doc, "9. Tier Design (Freemium → Enterprise)", 1)
    add_para(doc, "Five tiers. The freemium grid is designed to give consumers "
                  "enough value to stick around (3 saved articles, 3 searches per "
                  "month, 2 deep researches) but with the high-leverage features "
                  "(AOI alerts, URL analyses, embeds, exports) gated for paying "
                  "tiers. Enterprise unlocks team management, custom sources, "
                  "and API access.")

    add_heading(doc, "9.1 Tier Grid", 2)
    add_table(
        doc,
        ["Tier", "Price", "Saved", "Searches", "Deep Research", "AOI Alerts",
         "URL Analyses", "Exports"],
        [
            ["Anonymous", "Free", "0", "10/day IP", "0", "0", "0", "—"],
            ["Free (Freemium)", "€0", "3 total", "3/month", "2/month", "0", "0", "—"],
            ["Basic", "€10/mo", "25", "25/day", "10/month", "5", "5/month", "CSV"],
            ["Professional", "€20/mo", "unlimited", "50/day", "unlimited", "25", "unlimited", "CSV + PDF"],
            ["Enterprise", "Custom", "unlimited", "unlimited", "unlimited", "unlimited", "unlimited", "All + JSON/API"],
        ],
    )

    add_heading(doc, "9.2 Smart-Freemium Design Principles", 2)
    add_bullet(doc,
        "Free tier gives 3 of each high-value resource (saved articles / searches / "
        "deep researches). Three is enough to validate the workflow but not enough "
        "for sustained use — natural upgrade anchor.",
        bold_lead="Three-of-each rule — ")
    add_bullet(doc,
        "Anonymous users get IP-rate-limited browsing without sign-up friction. "
        "Drives organic discovery via search + share links.",
        bold_lead="Anonymous floor — ")
    add_bullet(doc,
        "URL analyses are NOT in the free tier — they're the platform's most "
        "expensive operation and its highest-perceived-value feature. Showing "
        "the surface but gating the action drives Basic upgrades.",
        bold_lead="URL analysis as upgrade carrot — ")
    add_bullet(doc,
        "AOI alerts (threshold-crossing email) need a real subscription to be "
        "valuable — gated to Basic+ accordingly. Journalists especially "
        "convert here for newsroom 'first to know' utility.",
        bold_lead="Alert gating — ")
    add_bullet(doc,
        "Embeds are Professional+ — they're a B2B utility (newsrooms, "
        "research firms) that anchors the Professional tier price.",
        bold_lead="Embeds for revenue — ")
    add_bullet(doc,
        "Enterprise unlocks team seats + custom RSS feed registration + "
        "API access + JSON exports. Targets ESG/Risk consulting firms.",
        bold_lead="Enterprise B2B — ")

    add_heading(doc, "9.3 Quota Enforcement", 2)
    add_para(doc, "QuotaService produces a structured 429 envelope when limits are "
                  "exceeded: {tier, limit, used, allowed, upgrade_url, label, "
                  "period, reset_at}. Frontend renders an UpgradeModal with the "
                  "specific upgrade CTA and the persona-appropriate copy.")

    add_page_break(doc)

    # ---- 10. Quality Gates ----
    add_heading(doc, "10. Quality Gates", 1)
    add_para(doc, "What's tested, where, and how regression prevention works.")

    add_heading(doc, "10.1 Backend Test Surface", 2)
    add_bullet(doc, "tests/api/ — FastAPI route tests with in-memory DB", bold_lead="API: ")
    add_bullet(doc, "tests/unit/domains/ — pure function tests", bold_lead="Domain: ")
    add_bullet(doc, "tests/integration/ — Celery + Redis integration", bold_lead="Integration: ")
    add_bullet(doc, "tests/adversarial/ — red-team probes against the verifier", bold_lead="Adversarial: ")
    add_bullet(doc, "153 tests in the core gate (test_company_routes / test_skills_registry / "
                    "test_agentic_skill_pin / test_numeric_grounding / test_multi_llm_verifier)", bold_lead="Core gate: ")

    add_heading(doc, "10.2 Frontend Test Surface", 2)
    add_bullet(doc, "src/frontend/src/__tests__/ — Vitest, 350 tests across components + pages", bold_lead="Vitest: ")
    add_bullet(doc, "src/frontend/e2e/ — Playwright, 60+ scenarios (accessibility, structured-failure, country-passport, deep-search-compare, quota-gate)", bold_lead="Playwright: ")
    add_bullet(doc, "TypeScript strict — 0 errors", bold_lead="Types: ")
    add_bullet(doc, "Next.js build succeeds with Suspense boundaries", bold_lead="Build: ")

    add_heading(doc, "10.3 Single-Source-of-Truth Pin Tests", 2)
    add_para(doc, "test_agentic_skill_pin.py asserts SKILLS_REGISTRY (Python) "
                  "matches both the chat prompt template's AVAILABLE ACTIONS list "
                  "and the frontend dispatcher's ChatActionType union. Any drift "
                  "fails CI and blocks merge — the LLM never gets to suggest a "
                  "skill the dispatcher can't execute.")

    add_heading(doc, "10.4 CI Pipeline", 2)
    add_bullet(doc, "GitHub Actions: lint + typecheck + backend pytest + frontend Vitest", bold_lead=".github/workflows/ci.yml: ")
    add_bullet(doc, "Adversarial probes run as a separate red-team gate", bold_lead="Probes: ")
    add_bullet(doc, "Cloud Build runs migrations + builds images + deploys on every push to main", bold_lead="Deploy gate: ")

    add_page_break(doc)

    # ---- 11. Known Limitations & Roadmap ----
    add_heading(doc, "11. Known Limitations & Roadmap", 1)
    add_para(doc, "Honest list of what's NOT yet at production-grade and the "
                  "scope of work to close each gap.")

    limits = [
        ("11.1 CDP + NZT live data sources",
         "CDP retired anonymous public CSV in 2024 (requires registration + API "
         "key). NZT moved their export to GraphQL. Both adapters return clean "
         "200 + 'data_source_unavailable' warnings; seed data covers the surface "
         "until live adapters land.",
         "Register for CDP API; build NZT GraphQL adapter. ~1-2 days each."),
        ("11.2 Full AR6 Atlas ingestion",
         "ProjectionsPanel uses seed-only data (20 countries × 3 SSPs × 3 horizons).",
         "Atlas ingestion adapter → 180 countries × 3 SSPs × 3 horizons + variability "
         "envelopes. ~1-2 days."),
        ("11.3 B4 NER entity grounding",
         "Numeric grounding catches number hallucinations; entity grounding (people, "
         "places, organisations) is deferred. Single-LLM entity hallucinations "
         "currently undetected.",
         "Add spaCy NER + cross-LLM entity-overlap check. ~2-3 days."),
        ("11.4 Cloud Tasks migration for background jobs",
         "FastAPI BackgroundTasks rely on min-instances=1 to complete reliably. "
         "Per-task work limits + retry handling would benefit from Cloud Tasks.",
         "Move adapter sync + AOI poll to Cloud Tasks queue with worker. ~2 days."),
        ("11.5 Per-company deduplication on live sync",
         "SBTi syncs can write duplicate company rows for multi-target companies "
         "without ticker/isin/lei. Migration 036 dedupes after the fact; "
         "upsert_company now has weak-fallback dedup but it runs per-row.",
         "Add a batch dedup pass at end of every adapter sync. ~0.5 day."),
        ("11.6 Frontend SSE / WebSocket for chat",
         "Chat answers stream as one chunk. UX would benefit from "
         "token-by-token streaming.",
         "Add SSE endpoint + frontend EventSource. ~1 day."),
    ]
    for title, status, work in limits:
        add_heading(doc, title, 2)
        add_kv(doc, "Status", status)
        add_kv(doc, "Closure", work)

    add_page_break(doc)

    # ---- 12. Production State Snapshot ----
    add_heading(doc, "12. Production State (snapshot at report time)", 1)

    add_heading(doc, "12.1 Live URLs", 2)
    add_kv(doc, "Frontend", "https://climatenews-frontend-srzwxdzmaq-ez.a.run.app")
    add_kv(doc, "API", "https://climatenews-api-srzwxdzmaq-ez.a.run.app")
    add_kv(doc, "GCP project", "climatenews-495412 (europe-west4)")
    add_kv(doc, "Repository", "github.com/eljaplacido/climatefactsai")

    add_heading(doc, "12.2 Data Inventory", 2)
    add_bullet(doc, "10,000+ rows (live SBTi data + 17 seed companies + duplicates pending cleanup)",
               bold_lead="companies: ")
    add_bullet(doc, "180 rows (20 countries × 3 SSPs × 3 horizons, AR6-aligned)",
               bold_lead="country_projections: ")
    add_bullet(doc, "200+ feeds covering UN-193 countries", bold_lead="rss_feed_registry: ")
    add_bullet(doc, "Multi-tier classification on every source", bold_lead="source_credibility_tiers: ")

    add_heading(doc, "12.3 Quality Gates (this deploy)", 2)
    add_bullet(doc, "153 tests pass (B3 + skills + numeric grounding + multi-LLM)", bold_lead="Backend: ")
    add_bullet(doc, "350 / 350 Vitest pass", bold_lead="Frontend: ")
    add_bullet(doc, "0 errors", bold_lead="TypeScript: ")
    add_bullet(doc, "All endpoints 200 except /api/companies/{ticker} which 500s "
                    "on a UUID-cast bug; fix queued in next deploy", bold_lead="Live: ")

    add_heading(doc, "12.4 Recent Migrations", 2)
    migs = [
        ("029", "corporate_disclosures (companies, company_climate_disclosures, company_claims)"),
        ("030", "drift_threshold_fits"),
        ("031", "url_analysis_failure_classification"),
        ("032", "aoi_subscriptions"),
        ("033", "source_credibility_tiers_expansion"),
        ("034", "corporate seed data (17 well-known public companies)"),
        ("035", "country_projections (180 rows, IPCC AR6)"),
        ("036", "dedupe_companies + partial-unique index on (name, country_code)"),
    ]
    for v, name in migs:
        add_bullet(doc, name, bold_lead=f"{v}: ")

    add_heading(doc, "12.5 Service-Level Indicators", 2)
    add_kv(doc, "API timeout", "1800s (Cloud Run max for HTTP)")
    add_kv(doc, "API min instances", "1 (warm for BackgroundTasks)")
    add_kv(doc, "API max instances", "5")
    add_kv(doc, "API concurrency", "80")
    add_kv(doc, "API memory", "1Gi")
    add_kv(doc, "Frontend min instances", "0")
    add_kv(doc, "Frontend max instances", "3")
    add_kv(doc, "Frontend memory", "256Mi")

    add_heading(doc, "12.6 Status Indicators (live checks)", 2)
    add_bullet(doc, "GET /api/skills → 200 with 11 skills", bold_lead="Agentic: ")
    add_bullet(doc, "GET /api/companies?limit=3 → 200 with company array", bold_lead="Corporate: ")
    add_bullet(doc, "GET /api/map/country/DE/projections → 200 with 3 SSP scenarios × 3 horizons",
               bold_lead="Projections: ")
    add_bullet(doc, "GET /api/methodology → 200 with bundled snapshot", bold_lead="Methodology: ")
    add_bullet(doc, "GET /api/companies/admin/sync/sbti → 200 with last-run status",
               bold_lead="Adapter status: ")

    add_callout(doc,
        f"Report generated {date.today().isoformat()} from a deployed snapshot. "
        "Numbers may drift as live SBTi syncs add rows and dedup migration runs.",
        color=MUTED)

    # ---- Write ----
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote: {OUT_PATH}")
    print(f"Size:  {OUT_PATH.stat().st_size:,} bytes")


if __name__ == "__main__":
    build()

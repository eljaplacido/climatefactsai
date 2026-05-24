"""
Generate the Climatefacts.ai improvement-plan synthesis as a proper .docx.

Output: docs/reports/Climatefacts-Improvement-Synthesis-2026-05-19.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# --- colours --------------------------------------------------------------
BRAND = RGBColor(0x0E, 0x5C, 0x4A)         # deep teal — headings
BRAND_2 = RGBColor(0x1A, 0x6B, 0x57)       # mid teal — h3
INK = RGBColor(0x11, 0x11, 0x11)
SUBTLE = RGBColor(0x55, 0x55, 0x55)
DONE = RGBColor(0x0A, 0x6F, 0x40)
OPEN = RGBColor(0x9A, 0x4D, 0x00)
WARN_BORDER = "C25450"
CALLOUT_BORDER = "0E5C4A"
HEADER_FILL = "E8F1EE"
CODE_FILL = "F2F2F2"


# --- helpers --------------------------------------------------------------
def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, hex_color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def set_paragraph_border(paragraph, side: str, hex_color: str, size: int = 24) -> None:
    """size is in eighths of a point (24 = 3pt)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bdr = OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), str(size))
    bdr.set(qn("w:space"), "4")
    bdr.set(qn("w:color"), hex_color)
    p_borders.append(bdr)
    p_pr.append(p_borders)


def configure_styles(doc: Document) -> None:
    """Set up base font, heading colours, table grid look."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    title = doc.styles["Title"]
    title.font.name = "Calibri Light"
    title.font.size = Pt(28)
    title.font.color.rgb = BRAND
    title.font.bold = True

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri Light"
    h1.font.size = Pt(18)
    h1.font.color.rgb = BRAND
    h1.font.bold = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.color.rgb = BRAND_2
    h2.font.bold = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.color.rgb = INK
    h3.font.bold = True


def add_para(doc, text: str, *, italic=False, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if italic:
        run.italic = True
    if bold:
        run.bold = True
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return p


def add_rich_para(doc, segments):
    """segments = list of (text, dict-of-style) e.g. ('Foo', {'bold': True})."""
    p = doc.add_paragraph()
    for text, style in segments:
        run = p.add_run(text)
        if style.get("bold"):
            run.bold = True
        if style.get("italic"):
            run.italic = True
        if style.get("mono"):
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        if style.get("color"):
            run.font.color.rgb = style["color"]
    return p


def add_bullets(doc, items):
    for item in items:
        if isinstance(item, str):
            doc.add_paragraph(item, style="List Bullet")
        else:
            p = doc.add_paragraph(style="List Bullet")
            for text, style in item:
                run = p.add_run(text)
                if style.get("bold"):
                    run.bold = True
                if style.get("italic"):
                    run.italic = True
                if style.get("mono"):
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc, text, *, kind="info"):
    """Coloured left-bordered box with shaded background."""
    p = doc.add_paragraph()
    border_color = WARN_BORDER if kind == "warn" else CALLOUT_BORDER
    fill_color = "FBF3F2" if kind == "warn" else "F4FAF8"
    set_paragraph_border(p, "left", border_color, size=32)
    set_paragraph_shading(p, fill_color)
    runs = text if isinstance(text, list) else [(text, {})]
    for t, style in runs:
        run = p.add_run(t)
        if style.get("bold"):
            run.bold = True
        if style.get("italic"):
            run.italic = True
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)


def add_table(doc, headers, rows, *, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    # header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = BRAND
        run.font.size = Pt(10)
        set_cell_bg(cell, HEADER_FILL)

    # body
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            if isinstance(val, list):
                for text, style in val:
                    run = para.add_run(text)
                    if style.get("bold"):
                        run.bold = True
                    if style.get("italic"):
                        run.italic = True
                    if style.get("mono"):
                        run.font.name = "Consolas"
                        run.font.size = Pt(9)
                    else:
                        run.font.size = Pt(10)
                    if style.get("color"):
                        run.font.color.rgb = style["color"]
            else:
                run = para.add_run(str(val))
                run.font.size = Pt(10)

    if col_widths:
        for row in table.rows:
            for c_idx, w in enumerate(col_widths):
                row.cells[c_idx].width = w
    return table


def add_code_block(doc, text: str):
    p = doc.add_paragraph()
    set_paragraph_shading(p, "F6F6F6")
    set_paragraph_border(p, "left", "DDDDDD", size=16)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = INK


def add_hr(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bdr = OxmlElement("w:bottom")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), "6")
    bdr.set(qn("w:space"), "1")
    bdr.set(qn("w:color"), "BBBBBB")
    p_borders.append(bdr)
    p_pr.append(p_borders)


# --- content --------------------------------------------------------------
def build():
    doc = Document()
    configure_styles(doc)

    # ---------- cover ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Improvement-Plan Synthesis & Impact-Radius Analysis")
    run.font.name = "Calibri Light"
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND
    run.bold = True

    sub = doc.add_paragraph()
    r = sub.add_run(
        "Cross-reading three external improvement plans against the live codebase — "
        "with file-level impact graphs ready for gitnexus visualisation"
    )
    r.font.size = Pt(13)
    r.font.color.rgb = SUBTLE
    r.italic = True

    meta = doc.add_paragraph()
    for line in [
        ("Date: ", True), ("19 May 2026", False), ("\n", False),
        ("Subject platform: ", True), ("Climatefacts.ai (CISU Regen)", False), ("\n", False),
        ("Repo state: ", True), ("branch main — HEAD 4699a70 | prior anchor 8d34329 (launch hardening) | rebrand 1b96d3a", False), ("\n", False),
        ("Inputs analysed:", True),
    ]:
        run = meta.add_run(line[0])
        run.bold = line[1]
        run.font.size = Pt(11)
    add_bullets(doc, [
        "Climatefacts-Strategic-Analysis-2026-05-18.md  (5-tier strategic / regulatory framing)",
        "Climate Platform Analysis and Improvement Plan.md  (greenwashing-NLP and corporate-disclosure focus)",
        "Climate Data Platform: Strategic Analysis & Roadmap to Best-in-Class (1).md  (FAIR + dMRV + blockchain provenance focus)",
    ])
    add_hr(doc)

    # ---------- 1. What changed since 2026-05-18 ----------
    doc.add_heading("1. What changed since the 2026-05-18 platform report", level=1)
    doc.add_paragraph(
        "Before mapping the three improvement plans onto the codebase, the picture must be re-anchored: "
        "two large commits landed after the previous report was filed, and they have already closed "
        "many of the items the plans flag."
    )
    add_table(doc,
        headers=["Commit", "Date", "What landed"],
        rows=[
            ["8d34329", "2026-05-18",
             "Launch-hardening mega-commit. RSS full-body extraction (BeautifulSoup + timeout + summary fallback) shipped in rss_adapter.py. is_synthetic = FALSE filter applied to 60+ user-facing SELECTs across 13 route files. Sustainability v2 with ND-GAIN wired in. Calibration is_preview tag when n_labels < 50. Hallucination-rates schema fix. Scheduler fail-closed + CSP header + /terms + /privacy. 196 country codes via migration 026. 658 articles enriched."],
            ["4699a70", "2026-05-19",
             "Tier-1 analytical remediations. Sustainability v2 now adds UNFCCC NDC as the 5th component (weights 0.30/0.25/0.20/0.15/0.10); year_spread confidence-band widening. Multi-LLM numeric grounding via optional callable in multi_llm_verifier.py; numbers_in_claim and numeric_grounded fields on CorroboratedClaim. spaCy NER hallucination entity check (graceful fallback). UI/UX: /search skeleton grid, <main> landmark, aria-label on icon-only buttons. 477 tests passing."],
        ],
    )
    add_callout(doc, [
        ("Honest re-grading after these two commits. ", {"bold": True}),
        ("Of the 11 Tier-A 'credibility-floor' remediations in the strategic-analysis report, ", {}),
        ("8 are now shipped ", {"bold": True}),
        ("and the audited composite has moved from ~3.6 toward an estimated ", {}),
        ("~4.0–4.15/5 ", {"bold": True}),
        ("(still well short of the 4.78 self-claim, but materially closer). The three reports were filed against the pre-launch code; this synthesis re-reads them against HEAD.", {}),
    ])

    # ---------- 2. Consensus / Divergence ----------
    doc.add_heading("2. Where the three reports agree — and where they don't", level=1)
    doc.add_paragraph(
        "Each report comes from a distinct vantage point. Their overlap is the highest-confidence backlog; "
        "their divergence reveals which strategic bet the platform should choose deliberately."
    )

    doc.add_heading("2.1 Three-way consensus (high-confidence backlog)", level=2)
    add_table(doc,
        headers=["Recommendation", "Strategic", "Greenwashing-NLP", "FAIR / dMRV", "Status in code (HEAD 4699a70)"],
        rows=[
            ["Replace 8-publisher venue whitelist with DB-backed tiers", "C2", "§5.3", "Gap 1",
             [("Still open", {"bold": True, "color": OPEN}), (" — bayesian_credibility.py:48 hard-codes KNOWN_VENUES", {})]],
            ["Agentic chat actions / tool-use protocol", "B2", "§7.1", "(implicit)",
             [("Still open", {"bold": True, "color": OPEN}), (" — no actions payload in chat_routes.py or conversation_engine.py", {})]],
            ["NER + numeric grounding for hallucination/verification", "B4", "§3.3, 5.3", "Gap 5",
             [("Done", {"bold": True, "color": DONE}), (" — 4699a70 shipped both (numeric_grounded + spaCy)", {})]],
            ["Sustainability composite v2 (consume all 6 adapters)", "B5", "§5.3", "Gap 1",
             [("Largely done", {"bold": True, "color": DONE}), (" — 5 of 6 components live (IRENA still unused; needs population normaliser)", {})]],
            ["Corporate-claim verification module (CDP/SBTi/NZT)", "B3", "§1, 6, 7", "Gap 7",
             [("Not started", {"bold": True, "color": OPEN}), (" — zero matches for company_climate/cdp/sbti in src/backend/", {})]],
            ["Transparent uncertainty / confidence intervals in UI", "5.2 (decomposed)", "§3.2", "Gap 5 + Tier 1 #3",
             [("Partial", {"bold": True, "color": OPEN}), (" — DecomposedConfidenceChart exists but is unused; gauge still single-dial", {})]],
            ["Article-trail join: UUIDs → titles/URLs", "5.3", "§6.1", "Tier 1 #1",
             [("Open", {"bold": True, "color": OPEN}), (" — /api/methodology/audit-trail/* still returns opaque UUIDs", {})]],
            ["Calibration: raise min_labels, surface is_preview", "A3", "§5.3", "Tier 1 #3",
             [("Done", {"bold": True, "color": DONE}), (" — STABLE_FIT_MIN=50 on RefitResult.as_dict", {})]],
            ["RSS full-body extraction", "A1", "§5.3", "(pre-req)",
             [("Done", {"bold": True, "color": DONE}), (" — rss_adapter.py:_parse_feed fetches + extracts", {})]],
            ["Frontend filter on is_synthetic = FALSE", "A2", "§5.3", "(pre-req)",
             [("Done", {"bold": True, "color": DONE}), (" — 60+ API SELECTs filter; user surfaces clean", {})]],
        ],
    )

    doc.add_heading("2.2 Where the three diverge (the strategic choice)", level=2)
    add_table(doc,
        headers=["Vector", "Strategic (5-tier)", "Greenwashing-NLP", "FAIR / dMRV"],
        rows=[
            ["What the platform is",
             "Regulatory-aligned trust infrastructure for European ECGT / AI Act window",
             "Algorithmic assurance engine grading public greenwashing",
             "Universal trust layer / data commons with cryptographic provenance"],
            ["Largest moat to build",
             "Corporate-claim verification + AI Act labelling",
             "Adversarial multi-LLM verifier with auditor-persona prompts",
             "Immutable audit ledger + FAIR metadata + federated mesh"],
            ["Differentiating signal",
             "Published methodology standard (CFRS)",
             "NLP detection of greenwashing patterns",
             "Blockchain anchoring via Hyperledger / CAD Trust"],
            ["Defining UX move",
             "Decomposed gauge + persistent trust ribbon + reproduce-button",
             "Active agentic chat that shows its reasoning",
             "Confidence-interval display + verification badges (✓ / ✓✓ / ✓✓✓) + change-provenance log"],
            ["Long-horizon expansion",
             "SKUs: Claim Substantiation + CSRD Companion",
             "Video pipeline for verified-fact dissemination",
             "Nature/biodiversity, dMRV certification, federated mesh, Global South PWA"],
        ],
    )

    add_callout(doc, [
        ("The three reports do not contradict each other — but they cannot all be the next sprint. ", {"bold": True}),
        ("The strategic-analysis report is fastest to monetise (ECGT window is September 2026). The FAIR/dMRV report is the largest ", {}),
        ("category-defining ", {"italic": True}),
        ("bet (universal trust layer). The greenwashing-NLP report is the most defensible ", {}),
        ("technical ", {"italic": True}),
        ("moat (adversarial verifier). Pick the order deliberately; do not blend the headlines.", {}),
    ], kind="warn")

    # ---------- 3. Re-anchored backlog ----------
    doc.add_heading("3. Re-anchored backlog (verified against HEAD 4699a70)", level=1)
    doc.add_paragraph(
        "The 16 recommendations below are the union of the three reports' suggestions, "
        "filtered to what is still actually open after 8d34329 and 4699a70. Detailed file-level impact follows in §4."
    )
    add_table(doc,
        headers=["#", "Recommendation", "State", "Sourced from"],
        rows=[
            ["1",  "Source-tier DB (Scimago + RetractionWatch + IFCN + MBFC)",       "OPEN", "All three"],
            ["2",  "Agentic chat actions[] protocol + dispatchers",                  "OPEN", "Strategic B2, Greenwashing §7.1"],
            ["3",  "Audit-trail UUID↔title denormalisation",                         "OPEN", "All three"],
            ["4",  "AI Act Article-50 labelling layer (badges + JSON-LD)",           "OPEN", "Strategic B1"],
            ["5",  "Decomposed credibility gauge + persistent trust ribbon",         "OPEN", "Strategic 5.1–5.2, FAIR Tier 1 #3"],
            ["6",  "Reproduce-this-result button + diff view",                       "OPEN", "Strategic 5.7"],
            ["7",  "Per-country claim ledger view",                                  "OPEN", "Strategic 5.4"],
            ["8",  "Publish 4.78 vs 3.6 audit gap on /methodology",                  "OPEN", "Strategic A4 / §9"],
            ["9",  "Adversarial multi-LLM prompts (auditor persona)",                "OPEN", "Greenwashing §3.3"],
            ["10", "Corporate-claim MVP (CDP + SBTi + NZT) + /companies/[ticker]",   "OPEN", "Strategic B3, Greenwashing §1–3, FAIR Tier 2 #7"],
            ["11", "Provenance for negative findings",                               "OPEN", "Strategic 6.4"],
            ["12", "Indicator confidence (Bayesian σ per source)",                   "OPEN", "Strategic 6.2, FAIR Tier 1 #3"],
            ["13", "Learned drift thresholds",                                       "OPEN", "Strategic 6.3"],
            ["14", "Cross-language hallucination — per-language embeddings",         "OPEN", "Strategic C3"],
            ["15", "Embed widget + QR fact-check cards",                             "OPEN", "Strategic C5, FAIR §V"],
            ["16", "FAIR metadata + immutable audit ledger (Hyperledger / dMRV)",    "OPEN", "FAIR Tier 1 #1–2 + Tier 4 #16"],
        ],
    )

    # ---------- 4. Deep dive ----------
    doc.add_heading("4. Deep dive: implementation + impact radius", level=1)
    doc.add_paragraph(
        "Each item below carries the same four-block structure: what to do, impact graph "
        "(files/modules touched, ready to drop into gitnexus), effects (capability, grade-axis lift, "
        "regulatory fit), risk + rollback. Items are ordered by the recommended sequence in §5."
    )

    deep_dives = build_deep_dives()
    for rec in deep_dives:
        doc.add_heading(rec["title"], level=2)
        add_rich_para(doc, [("What. ", {"bold": True}), (rec["what"], {})])

        doc.add_heading("Impact graph", level=3)
        add_table(doc, headers=["", ""], rows=rec["impact"], col_widths=[Cm(4.5), Cm(11.5)])

        if rec.get("dag"):
            add_code_block(doc, rec["dag"])

        doc.add_heading("Effects", level=3)
        add_bullets(doc, rec["effects"])

        doc.add_heading("Risk & rollback", level=3)
        doc.add_paragraph(rec["risk"])

    # ---------- 5. Sequenced plan ----------
    doc.add_heading("5. Sequenced 12-week plan", level=1)
    add_table(doc,
        headers=["Week", "Recommendations", "Rationale"],
        rows=[
            ["1",     "R3 + R5 + R8",            "Three low-risk, high-trust-signal moves; composite climbs from ~4.0 to ~4.2 with two days of front-end work."],
            ["2",     "R4 + R1 (behind flag)",   "R4 cashes the EU AI Act August window; R1 closes the audit's largest single gap."],
            ["3–4",   "R2",                      "Biggest UX uplift; nine action types ship together; chat_actions_log seeds telemetry."],
            ["5",     "R9 + R11",                "Strengthens the multi-LLM verifier; opens negative-finding provenance — both prerequisites for R10."],
            ["6–8",   "R10",                     "The single highest-leverage strategic bet. Ships before September 2026 ECGT enforcement."],
            ["9",     "R7 + R6 + R15",           "Three distribution / journalistic-utility moves that complement the corporate-claim MVP."],
            ["10",    "R12 + R13",               "Two metrology upgrades that complete the sustainability and drift honesty stories."],
            ["11–12", "R14 (foundation)",        "Starts the longest-build item while corporate-claim MVP stabilises. The moat."],
        ],
    )
    doc.add_paragraph("What is not in this 12-week plan but stays in the long-horizon list:")
    add_bullets(doc, [
        [("R16 (FAIR + immutable ledger). ", {"bold": True}), ("Sequence Q3 2026, once corporate claims create a population of high-stakes claims worth anchoring.", {})],
        [("Nature / TNFD module ", {"bold": True}), ("(FAIR Tier 2 #5). October 2026 ISSB exposure draft is the trigger.", {})],
        [("Scope-3 supply-chain module ", {"bold": True}), ("After corporate-claim MVP proves out.", {})],
        [("Scenario engine + financial translation ", {"bold": True}), ("Pairs with the Insurance/Risk Enterprise SKU.", {})],
        [("Mobile PWA + WCAG 2.2 AA hardening. ", {"bold": True}), ("Quarterly cadence.", {})],
        [("S5 httpOnly cookies, D1 migration schism, D2 url_analyses UUID FK ", {"bold": True}), ("— three large chunks that fit a quiet week each.", {})],
    ])

    # ---------- 6. Gitnexus playbook ----------
    doc.add_heading("6. Gitnexus playbook — how to read this report inside the tool", level=1)
    doc.add_paragraph(
        "The 16 impact-graph blocks above are written in the shape gitnexus expects: a primary edit, "
        "a set of fan-out files by category (schema / API / frontend / tests / ops), and an explicit risk class. "
        "When the gitnexus CLI / MCP is online, the following queries reproduce each block as a live, "
        "click-through graph."
    )

    doc.add_heading("6.1 Query primitives", level=2)
    add_code_block(doc,
        "# Trace from a symbol outward\n"
        "gitnexus graph trace --from path/to/file.py:symbol --depth N\n\n"
        "# Fan-in: who currently calls this function?\n"
        "gitnexus graph fan-in --to path/to/file.py:symbol --depth N\n\n"
        "# Emit: dump the dependency cone of a folder\n"
        "gitnexus graph emit --from path/to/folder --include backend,frontend,tests\n\n"
        "# Impact-radius: simulate edit + propagate breaks\n"
        "gitnexus impact simulate --edit path/to/file.py:symbol --propagation tests,types,callers\n\n"
        "# Diff a planned change against the current graph\n"
        "gitnexus diff plan --plan docs/reports/Climatefacts-Improvement-Synthesis-2026-05-19.docx#R10"
    )

    doc.add_heading("6.2 Suggested ordered runs against this report", level=2)
    add_table(doc,
        headers=["Rec", "Gitnexus query"],
        rows=[
            ["R1",  "gitnexus graph fan-in --to src/backend/app/domains/intelligence/bayesian_credibility.py:KNOWN_VENUES --depth 4"],
            ["R2",  "gitnexus graph trace --from src/backend/app/domains/intelligence/conversation_engine.py --include frontend --depth 5"],
            ["R3",  "gitnexus graph trace --from api/methodology_routes.py:get_audit_trail --depth 3"],
            ["R4",  "gitnexus impact simulate --edit src/backend/app/domains/intelligence/provenance.py --propagation frontend,types"],
            ["R5",  "gitnexus graph fan-in --to src/frontend/src/components/DecomposedConfidenceChart.tsx --depth 2"],
            ["R9",  "gitnexus graph trace --from src/backend/app/domains/intelligence/multi_llm_verifier.py:verify_claims --depth 3"],
            ["R10", "gitnexus graph emit --from src/backend/app/domains/content/indicators/ --pattern-as corporate"],
            ["R14", "gitnexus impact simulate --edit src/backend/app/domains/content/embedding_service.py:DIM --propagation migrations,adapters,api,frontend"],
            ["R16", "gitnexus graph emit --from src/backend/services/ --pattern ledger_service"],
        ],
    )

    doc.add_heading("6.3 Risk-class colour mapping in the gitnexus UI", level=2)
    add_table(doc,
        headers=["Class", "Heuristic", "UI colour"],
        rows=[
            ["LOW",    "Additive; no schema; ≤5 files touched",                                                 "Green"],
            ["MEDIUM", "Schema migration OR >5 files OR feature flag required",                                 "Amber"],
            ["HIGH",   "Re-indexes data, changes score distribution, or requires partner co-signature",         "Red"],
        ],
    )

    # ---------- 7. Effects summary ----------
    doc.add_heading("7. Effects summary — if the full 12-week plan ships", level=1)
    add_table(doc,
        headers=["Axis", "Today (HEAD 4699a70)", "Post Week 12", "Driver(s)"],
        rows=[
            ["Reliability",                          "~4.1",                            "4.7",      "R1 + R10"],
            ["Transparency",                         "~4.7",                            "4.95",     "R3 + R4 + R6 + R8 + R11"],
            ["Traceability",                         "~4.5",                            "4.85",     "R3 + R11 + R16 (foundation)"],
            ["Calibration",                          "~3.5 (math right, labels <50)",   "4.0–4.3",  "R2 telemetry + ops labelling"],
            ["Hallucination",                        "~4.0 (post NER + numeric)",       "4.6",      "R9 + R11 + R14 (foundation)"],
            ["Drift detection",                      "~4.0",                            "4.7",      "R13"],
            ["Sustainability",                       "~4.0 (v2 + NDC + year_spread)",   "4.6",      "R12"],
            ["Corporate-claim coverage",             "1.0",                             "4.0",      "R10"],
            ["Multilingual integrity",               "3.0",                             "3.8",      "R14 (full lift Q3)"],
            ["Regulatory readiness (ECGT + Art 50)", "3.5",                             "5.0",      "R4 + R10"],
            ["Source-tier rigour",                   "1.5",                             "4.0",      "R1"],
            ["Distribution / category position",     "2.0",                             "3.5",      "R7 + R15 + R8"],
            [[("Honest composite", {"bold": True})], [("~4.0 / 5", {"bold": True})], [("~4.6 / 5", {"bold": True})], "+0.6 in 12 weeks"],
        ],
    )

    # ---------- 8. Single most important thing ----------
    doc.add_heading("8. The single most important thing", level=1)
    doc.add_paragraph(
        "If only one move is taken from this synthesis, take this one: ship Week 1 (R3 + R5 + R8) immediately, "
        "and on the same day publish the audit gap on /methodology. The act of doing this is the most powerful "
        "evidence the platform exists for. It is the inverse of every greenwashing pattern. It is the embodiment "
        "of 'Trust as a Service'. It costs nothing but courage. Every other recommendation in this report compounds "
        "from that act."
    )
    doc.add_paragraph(
        "The second-most-important move is to sequence R10 (corporate-claim MVP) into the September 2026 ECGT "
        "enforcement window. Climatefacts.ai is structurally a verification engine for environmental claims; "
        "the European market has been handed regulation that requires exactly such an engine. The three external "
        "reports agree on this. The window is short. The work is doable in three focused weeks because every "
        "public dataset (CDP open, SBTi commitments, Net Zero Tracker) is already public, free, and well-shaped "
        "for the existing adapter framework."
    )
    doc.add_paragraph("Everything else can wait.")

    add_hr(doc)
    end = doc.add_paragraph()
    r = end.add_run(
        "Document end. Compiled 2026-05-19 from branch main @ 4699a70. "
        "Companion to docs/reports/Climatefacts-Platform-Report-2026-05-18 and the three improvement "
        "plans under docs/improvementplans/."
    )
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = SUBTLE

    return doc


def build_deep_dives():
    """Returns ordered list of recommendations with their impact graphs."""
    return [
        # R1
        dict(
            title="R1. Source-tier database (replace 8-publisher whitelist)",
            what=(
                "Replace the hard-coded KNOWN_VENUES set in bayesian_credibility.py:48 with a database-backed "
                "source_credibility_tiers table seeded from objective scientometric data. Tiered prior: "
                "Tier 1 = +30 (Scimago Q1 + IFCN-verified), Tier 2 = +15 (Q2 / mainstream press with correction policy), "
                "Tier 3 = +5 (Q3–Q4 / NGO with sourcing), Unknown = 0, Retraction-flagged = −30."
            ),
            impact=[
                [[("Primary edit", {"bold": True})], "src/backend/app/domains/intelligence/bayesian_credibility.py (replace constant lookup with DB read; add LRU cache)"],
                ["Schema",        "New migration 027_source_credibility_tiers.sql → source_credibility_tiers(source_id, domain, doi_prefix, tier, evidence_url, last_audited_at, retracted_count) + seed loader"],
                ["New module",    "src/backend/app/domains/trust/source_tier_service.py + scripts/seed_source_tiers.py (Scimago JR CSV + RetractionWatch API)"],
                ["Callers",       "url_analysis_routes.py (compute_research_prior), article_enrichment_service.py, provenance.py (records resolved tier)"],
                ["API surface",   "GET /api/methodology/source-tiers + GET /api/methodology/source-tier/by-domain"],
                ["Frontend",      "MethodologyDrawer.tsx (new section); SourceProfileCard.tsx badge; ArticleCard.tsx hover state"],
                ["Tests",         "Refit tests/api/test_bayesian_credibility.py; new tests/api/test_source_tier_service.py + fixture CSV"],
                ["Ops",           "Add to nightly Celery beat: refresh-retraction-flags (RetractionWatch RSS poll)"],
            ],
            dag=(
                "gitnexus query:\n"
                "  graph trace --from bayesian_credibility.py:KNOWN_VENUES --depth 4\n"
                "expected fan-out: 6 backend modules, 3 routes, 1 migration, 2 frontend components, 4 tests\n"
                "risk class:       MEDIUM (low blast, but the prior enters every reliability score)"
            ),
            effects=[
                "Reliability axis lift: 3.5 → 4.5 (closes the audit's largest single gap).",
                "Regulatory fit: directly satisfies ECGT Art 3 ('recognised certification schemes') and CSRD assurance-readiness.",
                "Composite grade: +0.15 to +0.20.",
            ],
            risk=(
                "Medium. The tier is multiplicative on every credibility score; a bad seed can shift the score "
                "distribution platform-wide. Rollback: env flag CLILENS_USE_SOURCE_TIER_DB=false falls back to the legacy whitelist. "
                "Ship behind the flag for one week of A/B before flipping."
            ),
        ),
        # R2
        dict(
            title="R2. Agentic chat actions[] protocol + client dispatchers",
            what=(
                "Backend appends an optional actions: [{type, params, label}] array to every /api/chat and "
                "/api/articles/{id}/ask response. Each action is a user-confirmed click (the LLM never acts directly). "
                "Minimum nine types: navigate, analyze_url, apply_search_filters, apply_map_filters, "
                "open_methodology_section, open_country, start_deep_search, bookmark_article, start_calibration_label."
            ),
            impact=[
                [[("Primary edit", {"bold": True})], "src/backend/app/domains/intelligence/conversation_engine.py — emit structured ChatActionSpec alongside the answer"],
                ["Schema",        "New table chat_actions_log(action_id, session_id, message_id, action_type, params, was_clicked, clicked_at)"],
                ["API",           "api/chat_routes.py pass-through; api/conversation_routes.py records click events"],
                ["Prompt registry","intelligence/prompts.py — new chat_synthesis_with_actions prompt (versioned, fingerprinted)"],
                ["Frontend",      "AgenticAssistant.tsx + new chatActionDispatcher.ts + 9 adapters in src/frontend/src/lib/chat-actions/"],
                ["Deep-linkable routes", "/map?country=…, /search?credibility=…&tag=…, /analyze?url=…, /deep-search?q=…, /methodology#sustainability"],
                ["Tests",         "9 new dispatcher tests + 1 contract test (Zod / Pydantic schema validation)"],
            ],
            dag=(
                "gitnexus query:\n"
                "  graph trace --from conversation_engine.py:converse --depth 5 --include frontend\n"
                "expected fan-out: 3 backend modules, 2 routes, 1 prompt, 12 frontend files, 11 tests\n"
                "risk class:       HIGH (LLM-driven UX state changes; needs user-confirm-click guardrail)"
            ),
            effects=[
                "UX axis lift: largest single move identified in the UI/UX audit. Reduces friction for institutional researchers.",
                "Telemetry side-effect: chat_actions_log is a high-quality dataset of which answers users actually act on — feeds calibration labels.",
                "Regulatory fit: when paired with R4 (AI Act badges), every dispatched action carries model + prompt version into the URL state.",
            ],
            risk=(
                "High. An LLM that produces malformed action params can navigate users to bad states. Mitigations: "
                "(1) every action passes through a Zod schema validator before being shown as a chip, (2) chips show the "
                "action label not raw JSON, (3) actions[] field is optional and feature-flagged off by default. "
                "Rollback: client ignores the field; the chat reverts to read-only."
            ),
        ),
        # R3
        dict(
            title="R3. Audit-trail UUID ↔ title denormalisation",
            what=(
                "The /api/methodology/audit-trail/{url-analysis|article|claim|deep-search}/{id} endpoints currently "
                "return source_article_ids as an opaque JSONB array of UUIDs. An external auditor cannot follow the "
                "chain without N additional API calls. Join articles in the audit-trail SQL and return "
                "[{article_id, title, url, source_name, source_tier, published_at}]."
            ),
            impact=[
                [[("Primary edit", {"bold": True})], "api/methodology_routes.py — four audit-trail endpoints; one shared _hydrate_source_articles(ids) helper"],
                ["Performance",   "Single CTE with WHERE article_id = ANY($1::uuid[]) — HNSW partial index already keys on article_id"],
                ["Frontend",      "MethodologyDrawer.tsx renders titles as clickable links; HallucinationBlock attaches per-source rate"],
                ["Tests",         "Refit tests/api/test_methodology_routes.py (already touched in 4699a70)"],
            ],
            dag=(
                "gitnexus query:\n"
                "  graph trace --from methodology_routes.py:get_audit_trail --depth 3\n"
                "expected fan-out: 1 backend module, 1 SQL change, 2 frontend components, 1 test file\n"
                "risk class:       LOW"
            ),
            effects=[
                "Transparency axis lift: 4.5 → 4.8 (closes the audit's 'opaque UUIDs' gap).",
                "Regulatory fit: CSRD assurance-readiness requires followable lineage; one-line schema response away.",
                "Composite grade: +0.05 to +0.10.",
            ],
            risk=(
                "Low. The endpoint is read-only and additive; no schema migration. Rollback: response field source_articles "
                "can be an array of objects without breaking existing UUID-only consumers if the legacy source_article_ids "
                "array is kept alongside."
            ),
        ),
        # R4
        dict(
            title="R4. AI Act Article-50 labelling layer",
            what=(
                "Add a persistent, machine-readable label to every AI-produced surface (chat answers, deep-search synthesis, "
                "URL-analysis verdicts, future videos). Three layers: visible badge ('AI-produced · v2026.05 · reproduce'), "
                "HTML attribute data-ai-generated='true', JSON-LD payload at footer with model, prompt_version, "
                "retrieval_strategy, timestamp, methodology_version."
            ),
            impact=[
                [[("Primary edit", {"bold": True})], "New component src/frontend/src/components/AIProvenanceBadge.tsx + jsonLdAi.ts serialiser"],
                ["Surfaces",     "ArticleDetailTabs.tsx, UrlAnalysisDetail, DeepSearchResult, AgenticAssistant message bubble, MethodologyDrawer"],
                ["Backend",      "provenance.py already records everything; add helper build_ai_provenance_payload(provenance_id)"],
                ["Public schema","docs/standard/ai-provenance.schema.json + /methodology#ai-transparency-standard publishes the schema"],
                ["Tests",        "Snapshot tests for badge + JSON-LD; one Cypress e2e that asserts the badge is present on every AI surface"],
            ],
            dag=(
                "gitnexus query:\n"
                "  graph fan-in --to provenance.py:build_payload --depth 2\n"
                "expected fan-out: 1 new component, 1 backend helper, 5 frontend surfaces, 1 spec doc\n"
                "risk class:       LOW-MEDIUM (additive)"
            ),
            effects=[
                "Regulatory fit: directly satisfies EU AI Act Article 50 (enforcement window August 2026). Most platforms will ship ugly retrofit banners; Climatefacts can ship this elegantly because it already has the data.",
                "Marketing surface: published schema becomes the Climatefacts AI Transparency Standard — first move toward C1 (CFRS publication).",
                "Composite grade: regulatory-readiness 3.0 → 4.5.",
            ],
            risk="Low. Visual change is intentional. Rollback: remove the component import; no data change.",
        ),
        # R5
        dict(
            title="R5. Decomposed credibility gauge + persistent trust ribbon",
            what=(
                "A single 0–100 gauge collapses information that ECGT, CSRD and the EU AI Act all require to be reported "
                "separately. Replace with the existing-but-unused DecomposedConfidenceChart as a four-bar view: "
                "Reliability · Agreement · Grounding · Calibration confidence. Add a persistent trust ribbon on article "
                "and analysis surfaces with four click-expandable icons."
            ),
            impact=[
                [[("Primary edit", {"bold": True})], "src/frontend/src/components/CredibilityGauge.tsx (deprecate single-dial primitive; gate via variant='decomposed')"],
                ["Already-existing", "DecomposedConfidenceChart.tsx — in repo, no callers; this recommendation wires it in"],
                ["Surfaces",      "ArticleCard, ArticleDetailTabs, UrlAnalysisDetail, MapCountryPanel, SearchResultCard"],
                ["New component", "TrustRibbon.tsx (icon row; provenance / calibration / hallucination / methodology version) + four click-expanding panels"],
                ["API",           "No backend change — data already on every relevant response"],
                ["Tests",         "Vitest snapshots; accessibility tests for keyboard-navigable disclosure panels"],
            ],
            dag=(
                "gitnexus query:\n"
                "  graph fan-in --to DecomposedConfidenceChart.tsx --depth 2\n"
                "expected: 0 current callers; this rec adds 5\n"
                "risk class: LOW"
            ),
            effects=[
                "UX axis lift: 4× information density on the same screen real estate.",
                "Regulatory fit: each ECGT-required dimension visible without expansion.",
                "Composite grade: transparency 4.7 → 4.9.",
            ],
            risk="Low. Frontend-only; behind config flag NEXT_PUBLIC_GAUGE_VARIANT=decomposed|single.",
        ),
        # R6
        dict(
            title="R6. Reproduce-this-result button",
            what=(
                "For any URL analysis or deep-search session, expose a 'Reproduce' button. Backend re-runs the analysis "
                "with the same prompt version + same retrieval strategy as recorded in claim_provenance and returns the "
                "diff (calibration shifted, new sources, new indicator data). The strongest possible demonstration that "
                "the platform is what it says it is."
            ),
            impact=[
                [[("Primary edit", {"bold": True})], "New module src/backend/app/domains/intelligence/reproducer.py — reads a provenance row, replays via the pinned prompt + retrieval, returns a diff struct"],
                ["API",          "POST /api/methodology/reproduce/{provenance_id} (admin or owner only) + GET /api/methodology/reproduce/{job_id}"],
                ["Celery",       "tasks/reproduce.py — async with WebSocket / SSE update channel"],
                ["Frontend",     "Button on UrlAnalysisDetail + DeepSearchResult; ReproductionDiff.tsx split-view"],
                ["Tests",        "Golden-fixture analysis with frozen prompt version (deterministic at T=0); regression test catching a methodology bump"],
            ],
            dag=(
                "gitnexus query:\n"
                "  graph emit --from reproducer.py --include provenance,prompts,deep_search_service\n"
                "risk class: MEDIUM (consumes LLM tokens; rate-limit per user)"
            ),
            effects=[
                "Transparency axis lift: 4.9 → 5.0 — this is the audit feature regulators will eventually require.",
                "Commercial: a documented reproducibility guarantee is a moat for the Enterprise tier.",
                "Composite grade: +0.05.",
            ],
            risk="Medium. LLM-tokens cost; per-user quota required. Rollback: remove the button; the endpoint can stay.",
        ),
        # R7
        dict(
            title="R7. Per-country claim ledger",
            what=(
                "Map drill-down today shows temperature anomalies + sustainability score + article count. Add a "
                "journalistically valuable view: chronological list of every climate-related claim made by or about that "
                "country in the last 12 months, with verdict, agreement score, and source. Turns the map from a "
                "comparison tool into a research tool."
            ),
            impact=[
                [[("Primary edit", {"bold": True})], "New endpoint GET /api/map/country/{cc}/claim-ledger?since=… in api/map_routes.py"],
                ["SQL",          "Query claims JOIN articles JOIN claim_provenance filtered by country_code + created_at >= now() - interval '12 months' + is_synthetic = FALSE"],
                ["Frontend",     "MapCountryPanel.tsx — new 'Claim ledger' tab; ClaimLedgerTable.tsx (sortable, filterable by verdict)"],
                ["Performance",  "Partial index idx_claims_country_date ON claims(country_code, created_at DESC) WHERE is_synthetic = FALSE (migration 028)"],
                ["Tests",        "Integration fixture with 5 claims across 2 countries; UI snapshot"],
            ],
            dag=None,
            effects=[
                "UX axis lift: high. The platform's most journalistically valuable view becomes the map's default drill-down.",
                "Distribution flywheel: this is the view that gets screenshotted in news pieces, which drives organic traffic.",
            ],
            risk="Low. New endpoint + new tab. Rollback: remove the tab; no data change.",
        ),
        # R8
        dict(
            title="R8. Publish the 4.78 / 3.6 audit gap on /methodology",
            what=(
                "The strongest trust signal the platform can send is to publish both the self-claimed grade and the "
                "audited grade alongside the remediation roadmap. The act of showing the gap inverts every greenwashing "
                "pattern."
            ),
            impact=[
                [[("Primary edit", {"bold": True})], "src/frontend/src/app/methodology/page.tsx — new 'Audited grade' section"],
                ["Optional API", "GET /api/methodology/grade-history returning array of (date, self_claim, audited, methodology_version, remediations)"],
                ["Static asset", "docs/audits/2026-05-18-*.md (already in repo) — the page links to them as primary evidence"],
            ],
            dag=None,
            effects=[
                "Marketing surface: zero-cost differentiation. Competitors will not copy this.",
                "Mission fit: directly embodies 'Trust as a Service'.",
            ],
            risk=(
                "None operational, some perception. The page is a content edit. Mitigate perception risk by framing as "
                "'Audit on rails — here's where we are, here's where we're going' rather than apologetically."
            ),
        ),
        # R9
        dict(
            title="R9. Adversarial multi-LLM prompts (auditor persona)",
            what=(
                "The Greenwashing-NLP report is correct that two LLMs running the same prompt collapse independence. "
                "4699a70 added numeric grounding but did not address the prompt-shared-bias failure mode. Add a second "
                "prompt template claim_extraction_auditor_persona with adversarial framing: 'You are a skeptical climate "
                "auditor. For each claim identify (i) vague modifiers, (ii) uncommitted future tense, "
                "(iii) baseline-decoupling, (iv) selective scope (1/2/3), (v) absence of validation body.' The secondary "
                "LLM uses this prompt; agreement now measures cross-model + cross-frame robustness."
            ),
            impact=[
                [[("Primary edit", {"bold": True})], "src/backend/app/domains/intelligence/prompts.py — new claim_extraction_auditor_persona v1.0 + fingerprint"],
                ["Verifier",    "multi_llm_verifier.py — optional secondary_prompt_name; records both prompts in raw_metadata.multi_llm_verification"],
                ["Extractor",   "anthropic_claim_extractor.py — accept prompt-name override"],
                ["Score",       "Greenwashing flags from the auditor persona become a new column raw_metadata.greenwashing_flags; surfaced on UrlAnalysisDetail"],
                ["Tests",       "5 fixture claims that match across both prompts; 5 that should diverge (vague-modifier, uncommitted-future, scope-omission, baseline-shifted, no-validation)"],
            ],
            dag=(
                "gitnexus query:\n"
                "  graph trace --from multi_llm_verifier.py:verify_claims --depth 3\n"
                "existing fan-out: 4 backend modules, 1 prompt, 2 routes, 6 tests\n"
                "risk class: MEDIUM (changes the score distribution of multi-LLM agreement)"
            ),
            effects=[
                "Multi-LLM axis lift: 3.5 → 4.5.",
                "Greenwashing detection: net-new capability — surfaces ECGT-prohibited claim patterns (vague modifiers etc.) directly in the audit trail.",
                "Composite grade: +0.10 to +0.15.",
            ],
            risk=(
                "Medium. Auditor-persona prompts are more likely to extract claims the primary missed — the verifier "
                "needs careful threshold tuning. Env flag CLILENS_AUDITOR_PERSONA=on for a one-week shadow period."
            ),
        ),
        # R10
        dict(
            title="R10. Corporate-claim verification MVP — /companies/[ticker]",
            what=(
                "The single largest strategic gap. All three reports name this. The platform verifies news; "
                "2026 regulation (ECGT, CSRD, IFRS S2) is overwhelmingly about verifying corporate climate claims. "
                "Minimum viable shape: a new entity type, three public datasets (CDP open + SBTi commitments + "
                "Net Zero Tracker), three verdict types (SBTi-validated net-zero target / scope 1+2+3 disclosure coverage / "
                "claim recency), a new route mirroring /country/[cc], and an ECGT-aligned warning surface for offset-based "
                "'climate neutral' product claims."
            ),
            impact=[
                [[("New domain", {"bold": True})], "src/backend/app/domains/content/corporate/ with schemas.py, repository.py, services.py, three adapters cdp_adapter.py, sbti_adapter.py, net_zero_tracker_adapter.py"],
                ["Schema",        "Migration 029_corporate_disclosures.sql — companies(...), company_climate_disclosures(...), company_claims(...)"],
                ["Reuse",         "indicators/base.py ABC pattern; idempotent upsert; same indicator_sync_logs audit table"],
                ["Intelligence",  "Extend multi_llm_verifier.py with corporate-claim mode; reuse R9 auditor-persona prompt; new prompt corporate_claim_extraction v1.0"],
                ["API",           "New router api/company_routes.py with GET /api/companies, GET /api/companies/{ticker}, GET /api/companies/{ticker}/claims, POST /api/companies/{ticker}/analyze-claim"],
                ["Frontend",      "/companies/[ticker]/page.tsx + /companies/page.tsx (index); mirror /map country panel layout"],
                ["Methodology",   "New section 'Corporate-claim methodology' with the new prompts + adapters"],
                ["Scheduler",     "Three new Cloud Scheduler triggers: CDP monthly, SBTi monthly, NZT weekly"],
                ["Tests",         "3 adapter test files + integration test for analyze-claim + end-to-end for the ticker page"],
            ],
            dag=(
                "gitnexus query:\n"
                "  graph emit --from corporate/ --include adapters,routes,frontend\n"
                "expected fan-out: ~15 new modules, 3 migrations, 4 routes, 2 frontend pages, 9 tests\n"
                "risk class: HIGH-BUILD (large surface), LOW-PRODUCTION (additive, no existing surfaces change)"
            ),
            effects=[
                "Strategic positioning: positions the platform for the September 2026 ECGT enforcement — the only window of regulatory tailwind in the 2026 calendar.",
                "Commercial: enables two new SKUs (Claim Substantiation, CSRD Companion).",
                "Composite grade: corporate-claim coverage axis 1.0 → 4.0.",
                "Effort: 2–3 weeks focused work with public datasets only — no new third-party APIs to negotiate.",
            ],
            risk=(
                "Medium. Greenfield code paths; large surface but no existing-surface coupling. Rollback: the new routes "
                "can be hidden behind NEXT_PUBLIC_ENABLE_COMPANIES=true. The backend modules can be left in place; no "
                "other code path is altered."
            ),
        ),
        # R11
        dict(
            title="R11. Provenance for negative findings",
            what=(
                "Today claim_provenance records what was produced. There's no record of what was looked for and not found. "
                "For a truth platform, negative space is information. When the multi-LLM verifier rejects a claim, when "
                "the hallucination check fires, when an indicator was unavailable for a country, when no contradicting "
                "source was found — each should be a row of its own type."
            ),
            impact=[
                [[("Schema", {"bold": True})], "Migration 030_negative_provenance.sql — widen claim_provenance.event_type enum to include claim_rejected, hallucination_flagged, indicator_missing, no_contradiction_found, numeric_grounding_failed"],
                ["Callers",  "multi_llm_verifier.py (write a row when corroboration fails), hallucination_detector.py (when score > threshold), sustainability_score.py (when an indicator is missing), deep_search_service.py (when synthesis abstains)"],
                ["API",      "New filter on /api/methodology/audit-trail/*: ?include_negative=true"],
                ["Frontend", "MethodologyDrawer shows 'What we looked for and didn't find' section when negative rows exist"],
            ],
            dag=None,
            effects=[
                "Transparency: complete provenance graph — the audit trail now answers 'is this absence of evidence informative?'",
                "Calibration input: negative events become training signal for calibration_labels (correct rejections improve the calibrated reliability score).",
            ],
            risk="Low. Pure schema widening + writes. Reads degrade gracefully if the new column types are absent.",
        ),
        # R12
        dict(
            title="R12. Indicator confidence — per-source σ",
            what=(
                "Each country_indicators row knows when it was ingested but not how confident the platform is in the "
                "upstream number itself. OWID per-capita emissions for Eritrea 2023 is plausibly ±30%; for Germany 2023 "
                "plausibly ±2%. The sustainability composite treats both as point estimates. Add an uncertainty_sigma "
                "column; propagate through the composite using standard error propagation."
            ),
            impact=[
                [[("Schema", {"bold": True})], "Migration 031_indicator_uncertainty.sql — country_indicators.uncertainty_sigma NUMERIC, country_indicators.distribution_shape VARCHAR (default normal)"],
                ["Adapters", "Each of the six adapters now writes the upstream's stated uncertainty (Climate TRACE publishes method-uncertainty per sector; OWID 1–5 quality tiers; CAT band-half-width; ND-GAIN confidence intervals)"],
                ["Composite","sustainability_score.py — propagate uncertainty via sigma_composite² = Σwᵢ² * σᵢ²; replaces the heuristic confidence_band with a defensible computed value"],
                ["API",      "Each indicator value now ships with {value, uncertainty, distribution}"],
                ["UI",       "Decomposed gauge (R5) renders error bars on each bar — ties R5 and R12 together"],
            ],
            dag=None,
            effects=[
                "Sustainability axis lift: 3.3 → 4.3.",
                "Equity correction: low-emitter countries with high-uncertainty data are no longer misleadingly shown as 'scoring 30/100' with a tight band.",
                "Composite grade: +0.05 to +0.10.",
            ],
            risk=(
                "Medium. Re-grades every sustainability score in production. Rollback: leave the column nullable; "
                "the composite falls back to the v2 fixed bands when uncertainty_sigma IS NULL."
            ),
        ),
        # R13
        dict(
            title="R13. Learned drift thresholds",
            what=(
                "KL drift thresholds (0.10 / 0.25 / 0.50 nats) are hard-coded in drift_detection.py:134. Collect 60 days "
                "of post-launch baseline data, fit a Gaussian to each drift signal's null distribution, set thresholds "
                "at 2σ / 3σ / 4σ. Publish the learned thresholds with their fit windows on /methodology."
            ),
            impact=[
                [[("Primary edit", {"bold": True})], "drift_detection.py — new fit_thresholds(signal, window_days) + persist into a new table drift_threshold_fits"],
                ["Celery", "New task tasks/drift_threshold_refit.py — weekly"],
                ["API",    "/api/methodology/drift-thresholds — surface the current fits + windows"],
            ],
            dag=None,
            effects=[
                "Drift axis lift: 4.0 → 4.7. Avoids the failure mode where a real drift event looks like baseline noise during a too-short baseline window.",
                "Publishable as metrology evidence — strongest data-quality signal a platform can put on its methodology page.",
            ],
            risk="Low. Replaces three constants with a fit. Constants remain as fallback.",
        ),
        # R14
        dict(
            title="R14. Cross-language hallucination — per-language embeddings + claim clustering",
            what=(
                "The Strategic-analysis report's original-research moat. The hardest problem in climate misinformation "
                "in 2026 is multilingual narrative laundering — a denialist claim debunked in English keeps spreading "
                "in Portuguese, Polish, Indonesian. Three sub-projects: (i) per-language embeddings — replace OpenAI "
                "ada-002 with bge-m3 or multilingual-e5-large; (ii) cross-language claim clustering — hash-key mapping "
                "the same proposition across languages; (iii) narrative-tracking dashboard — debunked claim's lifecycle."
            ),
            impact=[
                [[("Embedding swap", {"bold": True})], "src/backend/app/domains/content/embedding_service.py — provider-agnostic interface; new HNSW index for new dimensionality; migration 032_multilingual_embeddings.sql"],
                ["New domain", "src/backend/app/domains/intelligence/narrative_clustering.py — claim-cluster hash + cluster-membership table"],
                ["Schema",     "Migration 033_narrative_clusters.sql — narrative_clusters(cluster_id, canonical_text, language_count, member_count, first_seen_at) + claim_cluster_membership(claim_id, cluster_id, language_code, similarity)"],
                ["API",        "/api/narratives, /api/narratives/{cluster_id} (lifecycle), /api/narratives/{cluster_id}/timeline"],
                ["Frontend",   "/narratives — explorer of cross-language narrative clusters"],
                ["Backfill",   "One-off Celery job to re-embed historical articles (~5 hours on local corpus); cluster build (~1 hour)"],
            ],
            dag=(
                "gitnexus query:\n"
                "  graph emit --from narrative_clustering.py --include embeddings,migrations,routes\n"
                "expected fan-out: ~8 new modules, 2 migrations, 3 routes, 1 frontend route, 6 tests\n"
                "risk class: HIGH-BUILD (most ambitious item; original research)"
            ),
            effects=[
                "Category position: this is the moat. No other platform has it.",
                "Multilingual integrity axis: 3.0 → 4.7.",
                "Mission fit: directly delivers 'anyone, anywhere, in any language' with a structural capability no competitor will match in <12 months.",
            ],
            risk=(
                "High. Embedding model swap re-indexes the corpus; the cluster identity stays language-agnostic which is "
                "structurally new. Rollback: keep both embedding indexes for one quarter; route /narratives hidden behind a flag."
            ),
        ),
        # R15
        dict(
            title="R15. Embed widget + QR fact-check cards",
            what=(
                "The light version of the deferred Phase-9 video pipeline. Auto-generate one-claim fact-check cards "
                "(1080×1080 + 1080×1920) from any URL-analysis result, with verdict + two-source citation + methodology "
                "link + QR code back to the full audit trail. Plus an <iframe> embed widget for newsrooms; plus a "
                "Slack/Teams app."
            ),
            impact=[
                [[("New module", {"bold": True})], "api/factcard_routes.py — reuse og_image_routes.py's Pillow/Resvg pattern; add square + vertical variants + QR code"],
                ["Reuse",        "api/og_image_routes.py (already generates social cards); infographic_generator.py"],
                ["Embed widget", "src/frontend/src/app/embed/[analysis_id]/page.tsx — minimal iframe-friendly layout; <iframe> sandbox advice in share modal"],
                ["QR",           "qrcode library; URL pattern https://climatefacts.ai/a/{short_id}?utm=fact-card"],
                ["Slack/Teams",  "New service src/backend/services/slack_integration_service/ — uses existing /api/url-analysis"],
            ],
            dag=None,
            effects=[
                "Distribution: each shared card is a viral surface that links back to the platform.",
                "80% of the impact of the video pipeline at 20% of the engineering and zero of the social-platform-API risk.",
            ],
            risk="Low. Additive routes; nothing else changes.",
        ),
        # R16
        dict(
            title="R16. FAIR metadata + immutable audit ledger",
            what=(
                "The FAIR/dMRV report's Tier-1 anchor. Two sub-components: (i) FAIR metadata layer — persistent identifier "
                "per dataset (DOI-style), rich metadata aligned to F1–F4/A1/I1–I3/R1, registered with data.europa.eu / "
                "re3data.org, a FAIR-compliance score published per dataset; (ii) immutable ledger — anchor critical "
                "datasets (emissions inventories, carbon credits, NDC reports) to a Hyperledger Fabric consortium chain. "
                "Each anchoring event records data hash + timestamp + submitting entity + methodology applied."
            ),
            impact=[
                [[("FAIR module", {"bold": True})], "src/backend/app/domains/trust/fair_metadata.py + per-dataset YAML manifests under infrastructure/fair/datasets/"],
                ["Ledger module", "src/backend/services/ledger_service/ — Hyperledger Fabric SDK + Merkle-tree anchoring + verification CLI"],
                ["Schema",        "Migration 034_ledger_anchors.sql — ledger_anchors(anchor_id, dataset_id, data_hash_sha256, merkle_root, fabric_txid, anchored_at, methodology_version)"],
                ["API",           "/api/methodology/fair (compliance scores), /api/methodology/anchor/{id} (proof retrieval), /api/methodology/verify?dataset_id=…"],
                ["Frontend",      "Verification badge tier 3 (✓✓✓ cryptographically anchored) on every relevant surface"],
                ["Ops",           "Hyperledger network — consortium with partners (preferred) or single-org local for MVP; CAD Trust integration as fallback if Article-6.2 ITMOs become relevant"],
            ],
            dag=(
                "gitnexus query:\n"
                "  graph emit --from ledger_service/ --include all\n"
                "expected fan-out: ~12 new modules, 2 migrations, 4 routes, 1 service folder, 8 tests, 1 ops doc\n"
                "risk class: HIGH-BUILD + HIGH-OPS (needs partner co-signature)"
            ),
            effects=[
                "Category position: this defines the 'universal trust layer' the FAIR/dMRV report describes. Sets the platform up to be the infrastructure on which others build, not a product that competes with them.",
                "Regulatory fit: directly aligned with UNFCCC Climate Data Hub (Feb 2026), CAD Trust, Gold Standard dMRV pilot (running to October 2026).",
                "Commercial: the white-label / Enterprise SKU now has a defensible technical moat, not just a brand one.",
            ],
            risk=(
                "High. Most operationally complex item in the backlog. Sequence late — only after R10 creates a "
                "population of high-stakes claims worth anchoring. Rollback: ledger writes are write-through; reads "
                "fall back to DB-only audit trail with no functional loss."
            ),
        ),
    ]


# --- entry point ----------------------------------------------------------
def main():
    out_path = Path("docs/reports/Climatefacts-Improvement-Synthesis-2026-05-19.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(out_path)
    print(f"wrote {out_path} ({out_path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
